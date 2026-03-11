#!/usr/bin/env python3
"""
Интерфейс спутникового библиотекаря: интеграция Confluence с единой KB.
Проверка подключения, выбор пространства, синхронизация страниц в confluence_docs, перезагрузка KB.
"""
import os
import re
import sys
from pathlib import Path

# Корень проекта
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))
os.environ.setdefault("PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION", "python")

# Подгрузка config.env (CONFLUENCE_URL, CONFLUENCE_TOKEN)
_config_env = project_root / "config.env"
if _config_env.exists():
    with open(_config_env, "r") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip("'\""))

import io
import streamlit as st

from kb_billing.rag.confluence_client import ConfluenceClient


def _parse_spec_table(text: str):
    """
    Извлечь из ответа ассистента таблицу спецификации (наименование — N шт. или наименование — N).
    Возвращает (заголовок_или_None, list of (наименование, количество)) или (None, []) если не похоже на спецификацию.
    """
    if not text:
        return None, []
    lines = text.replace("\r\n", "\n").split("\n")
    # Строка с количеством в конце: "что угодно — 2 шт." (жадный .+ чтобы при "X — Y — 2 шт." взять "X — Y")
    row_pat = re.compile(r"^(.+)\s*[—–\-]\s*(\d+)\s*(?:шт\.?\s*)?$", re.IGNORECASE)
    title = None
    rows = []
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if "Скопируйте блок ниже" in line or "сохраните как" in line.lower():
            continue
        # Точное совпадение: "Наименование — 2 шт." или "Терминал — 1"
        row_match = row_pat.match(line)
        if row_match:
            name_part = row_match.group(1).strip()
            qty = int(row_match.group(2))
            # Не считать заголовком таблицы
            if not rows and name_part.lower() in ("наименование", "название") and qty == 1 and "кол" not in line.lower():
                continue
            rows.append((name_part, qty))
            continue
        # Разбивка по первому тире, если правая часть — число (или число + шт)
        for sep in (" — ", " – ", " - ", "\t"):
            if sep in line:
                parts = line.split(sep, 1)
                if len(parts) == 2:
                    name_part = parts[0].strip()
                    right = parts[1].strip()
                    num_match = re.match(r"^(\d+)\s*(?:шт\.?\s*)?$", right, re.IGNORECASE)
                    if num_match:
                        rows.append((name_part, int(num_match.group(1))))
                        break
                    # В наименовании может быть ещё тире: "X — Y — 2 шт." — ищем число в конце
                    num_trailing = re.search(r"(\d+)\s*(?:шт\.?\s*)?$", right, re.IGNORECASE)
                    if num_trailing:
                        rows.append((name_part + sep + right[: num_trailing.start()].strip().rstrip("—–-").strip(), int(num_trailing.group(1))))
                        break
                break
        else:
            # Не строка таблицы — возможно заголовок (судно, хаб, дата)
            if title is None and not rows and len(line) > 2 and not re.search(r"[—–\-]\s*\d+\s*(шт\.?\s*)?$", line, re.IGNORECASE):
                title = line
    return (title, rows) if rows else (None, [])


def _build_spec_docx(spec_title: str, spec_rows: list) -> bytes:
    """Собрать документ Word (.docx) со спецификацией: заголовок + таблица Наименование/Количество."""
    from docx import Document
    doc = Document()
    if spec_title:
        p = doc.add_paragraph(spec_title)
        p.runs[0].bold = True
    table = doc.add_table(rows=len(spec_rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Наименование"
    hdr[1].text = "Количество"
    for r, (name, qty) in enumerate(spec_rows, start=1):
        row = table.rows[r].cells
        row[0].text = str(name)
        row[1].text = str(qty)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


try:
    from kb_billing.rag.confluence_kb_generator import ConfluenceKBGenerator
    HAS_CONFLUENCE_GENERATOR = True
except ImportError as e:
    HAS_CONFLUENCE_GENERATOR = False
    ConfluenceKBGenerator = None

try:
    from kb_billing.rag import satellite_chat_db
    HAS_CHAT_DB = True
except ImportError:
    satellite_chat_db = None
    HAS_CHAT_DB = False


def _get_client(url: str = "", token: str = "") -> ConfluenceClient:
    url = url or os.getenv("CONFLUENCE_URL", "")
    token = token or os.getenv("CONFLUENCE_TOKEN", "")
    return ConfluenceClient(base_url=url or None, token=token or None)


def show_confluence_librarian_tab():
    """Закладка «Спутниковый ассистент» — интеграция Confluence с KB."""
    st.header("🛰️ Спутниковый ассистент — Confluence и KB")
    if not HAS_CONFLUENCE_GENERATOR:
        st.error(
            "Модуль `confluence_kb_generator` не найден. Синхронизируйте папку `kb_billing/rag/` на сервер "
            "(в т.ч. файл `confluence_kb_generator.py`) и установите зависимости: `pip install beautifulsoup4 requests`."
        )
        st.info("Проверка подключения к Confluence ниже должна работать.")
    st.markdown("""
    **Синхронизация документации и схем из Confluence в единую базу знаний:**
    - 🔗 Проверка подключения к Confluence (docs.steccom.ru)
    - 📂 Список пространств и выбор пространства для выгрузки
    - 📥 Синхронизация страниц в формат KB (сохранение в `confluence_docs/`)
    - 🔄 Обновление векторной KB (перезагрузка в Qdrant)
    """)

    tab_chat, tab_retrieval, tab_editor, tab_sync = st.tabs(
        [
            "💬 Ответы (инженер / абонент)",
            "🔎 Контекст (RAG поиск)",
            "✍️ Редактор KB",
            "🔄 Confluence / Qdrant",
        ]
    )

    confluence_docs_dir = project_root / "kb_billing" / "confluence_docs"

    def _load_json_list(path: Path):
        import json as _json
        if not path.exists():
            return []
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = _json.load(f)
            if data is None:
                return []
            return data if isinstance(data, list) else [data]
        except Exception:
            return []

    def _save_json_list(path: Path, data):
        import json as _json
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            _json.dump(data, f, ensure_ascii=False, indent=2)

    def _render_retrieval_test():
        st.subheader("Проверка релевантности (тест поиска по KB)")
        st.markdown(
            "Задайте вопрос — ниже будут показаны **какие именно чанки каких документов** найдены: документ, секция (или вложение), релевантность (score) и фрагмент текста. Если вставите ссылку на вложение Confluence (download/attachments/...), поиск выполнится по странице, к которой прикреплён файл."
        )
        relevance_question = st.text_area(
            "Вопрос для проверки поиска",
            value="",
            height=80,
            placeholder="Например: как настроить Kingsat P8? iDirect X5; SeaTel коды ошибок; порядок действий при некорректной работе...",
            key="relevance_question",
        )
        relevance_limit = st.slider("Сколько чанков показать", min_value=1, max_value=20, value=10, key="relevance_limit")
        snippet_len = st.slider("Длина фрагмента в превью (символов)", min_value=200, max_value=2000, value=600, key="snippet_len")
        if st.button("🔍 Найти в документах", key="relevance_search_btn"):
            if not relevance_question.strip():
                st.warning("Введите вопрос.")
                return
            with st.spinner("Поиск по векторной KB..."):
                try:
                    import re
                    from kb_billing.rag.rag_assistant import RAGAssistant
                    assistant = RAGAssistant()
                    q = relevance_question.strip()
                    docs = []
                    m = re.search(r"/download/attachments/(\d+)(?:/|$|\?)", q, re.IGNORECASE)
                    if m:
                        page_id = m.group(1)
                        docs = assistant.search_semantic(
                            "схема изображение вложение описание",
                            content_type="confluence_section",
                            limit=relevance_limit,
                            page_id=page_id,
                        )
                        if not docs:
                            docs = assistant.get_confluence_chunks_by_page_id(page_id, limit=relevance_limit)
                    if not docs:
                        docs = assistant.search_semantic(
                            q,
                            content_type="confluence_section",
                            limit=relevance_limit,
                        )
                    if not docs:
                        docs = assistant.search_semantic(
                            q,
                            content_type="confluence_doc",
                            limit=relevance_limit,
                        )
                    if not docs:
                        in_qdrant = assistant.get_confluence_page_ids()
                        if not in_qdrant:
                            hint = (
                                "В Qdrant **нет ни одного чанка Confluence**. Поиск в этой вкладке идёт только по спутниковой KB (Confluence). "
                                "Синхронизируйте страницы во вкладке «Confluence / Qdrant» (пространство или конкретные URL) и нажмите **«Перезагрузить в Qdrant только Confluence»**."
                            )
                        elif m:
                            page_id = m.group(1)
                            if page_id not in in_qdrant:
                                hint = (
                                    f"По странице **page_id={page_id}** чанков нет. В коллекции есть страницы: {', '.join(in_qdrant[:20])}"
                                    + (" …" if len(in_qdrant) > 20 else "")
                                    + ". Синхронизируйте нужную страницу по URL и перезагрузите Confluence в Qdrant."
                                )
                            else:
                                hint = f"Чанки со страницы {page_id} есть, но по запросу ничего не вернулось. Увеличьте «Сколько чанков показать» или попробуйте другой запрос."
                        else:
                            hint = (
                                f"По запросу чанков не найдено. В коллекции есть документы со страниц: {', '.join(in_qdrant[:15])}"
                                + (" …" if len(in_qdrant) > 15 else "")
                                + ". Попробуйте слова: **VSAT**, инструкция, схема, контроллер, антенна, модем, SeaTel, Intellian, C-COM."
                            )
                        st.info("По запросу ничего не найдено. " + hint)
                        return
                    st.success(f"Найдено чанков: **{len(docs)}**")
                    st.caption("Ниже: для каждого чанка — **документ**, **секция/вложение**, **релевантность**, фрагмент текста.")
                    for i, d in enumerate(docs, 1):
                        score = d.get("similarity", 0)
                        title = d.get("title", "—")
                        section_title = d.get("section_title", "") or "(без названия секции)"
                        url = d.get("source_url", "")
                        content = d.get("content") or ""
                        short = content[: snippet_len] + ("..." if len(content) > snippet_len else "")
                        chunk_key = f"show_full_chunk_{i}"
                        expanded = st.session_state.get(chunk_key, False)
                        with st.expander(f"**Чанк {i}** · Релевантность: **{score:.2%}**", expanded=expanded):
                            st.markdown("**Документ:** " + title)
                            st.markdown("**Секция / чанк:** " + section_title)
                            st.markdown("**Релевантность (score):** " + f"{score:.2%}")
                            if url:
                                st.markdown(f"📎 [Открыть в Confluence]({url})")
                            st.markdown("**Фрагмент текста:**")
                            st.text(short)
                            if len(content) > snippet_len:
                                show_full = st.checkbox(
                                    "Показать полный текст чанка",
                                    value=expanded,
                                    key=chunk_key,
                                )
                                if show_full:
                                    st.text(content)
                except Exception as e:
                    st.error(str(e))
                    import traceback
                    st.code(traceback.format_exc())

    def _render_manual_kb_editor():
        st.subheader("✍️ Ручные заметки для спутниковой KB")
        st.caption(
            "Здесь можно добавлять/редактировать заметки напрямую. Они сохраняются в `kb_billing/confluence_docs/manual_notes.json` "
            "и будут индексироваться как обычные документы Confluence (domain=satellite) после «Перезагрузить KB в Qdrant»."
        )

        manual_path = confluence_docs_dir / "manual_notes.json"
        docs = _load_json_list(manual_path)

        st.markdown("#### Быстро добавить заметку")
        col_a, col_b = st.columns(2)
        with col_a:
            doc_title = st.text_input("Заголовок документа", value="", key="manual_doc_title")
            section_title = st.text_input("Заголовок секции", value="", key="manual_section_title")
            scope_raw = st.text_input("Scope (через запятую)", value="general", key="manual_scope")
        with col_b:
            source_url = st.text_input("Источник (URL, опционально)", value="", key="manual_source_url")
            section_text = st.text_area("Текст секции", value="", height=160, key="manual_section_text")

        if st.button("💾 Сохранить заметку в manual_notes.json", key="manual_save_btn"):
            if not doc_title.strip() or not section_text.strip():
                st.error("Нужны как минимум: заголовок документа и текст секции.")
            else:
                import hashlib
                from datetime import datetime, timezone

                title_norm = doc_title.strip()
                now = datetime.now(timezone.utc).isoformat()
                existing = next((d for d in docs if (d.get("title") or "").strip() == title_norm), None)
                if existing:
                    page_id = (existing.get("source") or {}).get("page_id") or f"manual:{hashlib.md5(title_norm.encode('utf-8')).hexdigest()[:10]}"
                    existing.setdefault("source", {})
                    existing["source"]["page_id"] = page_id
                    if source_url.strip():
                        existing["source"]["url"] = source_url.strip()
                    existing["source"]["last_updated"] = now
                    existing.setdefault("content", [])
                    existing["content"].append(
                        {
                            "title": (section_title.strip() or "Заметка"),
                            "text": section_text.strip(),
                            "subsections": [],
                        }
                    )
                    sc = [s.strip() for s in (scope_raw or "").split(",") if s.strip()]
                    if sc:
                        existing["scope"] = sc
                else:
                    page_id = f"manual:{hashlib.md5(title_norm.encode('utf-8')).hexdigest()[:10]}"
                    sc = [s.strip() for s in (scope_raw or "").split(",") if s.strip()] or ["general"]
                    docs.append(
                        {
                            "title": title_norm,
                            "content": [
                                {
                                    "title": (section_title.strip() or "Заметка"),
                                    "text": section_text.strip(),
                                    "subsections": [],
                                }
                            ],
                            "source": {
                                "page_id": page_id,
                                "url": source_url.strip(),
                                "last_updated": now,
                            },
                            "scope": sc,
                        }
                    )
                _save_json_list(manual_path, docs)
                st.success("Сохранено. Теперь нажмите «Перезагрузить KB в Qdrant» (вкладка «Confluence / Qdrant»).")
                st.rerun()

        st.markdown("---")
        st.markdown("#### Редактирование JSON (полный контроль)")
        import json as _json
        raw = _json.dumps(docs, ensure_ascii=False, indent=2)
        edited = st.text_area("manual_notes.json", value=raw, height=320, key="manual_raw_json")
        col_v, col_s = st.columns(2)
        with col_v:
            if st.button("✅ Проверить JSON", key="manual_validate_json_btn"):
                try:
                    _json.loads(edited)
                    st.success("JSON валиден.")
                except Exception as e:
                    st.error(f"Ошибка JSON: {e}")
        with col_s:
            if st.button("💾 Сохранить JSON", type="primary", key="manual_save_json_btn"):
                try:
                    parsed = _json.loads(edited)
                    if parsed is None:
                        parsed = []
                    if not isinstance(parsed, list):
                        parsed = [parsed]
                    _save_json_list(manual_path, parsed)
                    st.success("Сохранено. Теперь перезагрузите KB в Qdrant.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Не удалось сохранить: {e}")

    def _render_confluence_docs_list():
        st.subheader("📋 Документы Confluence в KB")
        st.caption(
            "Список документов из `kb_billing/confluence_docs/`. "
            "Можно обновить документ из Confluence или пометить устаревшим (тогда он не попадёт в поиск до «Перезагрузить KB»)."
        )
        if not confluence_docs_dir.exists():
            st.caption("Папка `kb_billing/confluence_docs/` пока пуста — после первой синхронизации здесь появятся документы.")
            return

        outdated_set = set()
        outdated_file = confluence_docs_dir / "outdated.txt"
        if outdated_file.exists():
            try:
                with open(outdated_file, "r", encoding="utf-8") as f:
                    outdated_set = {line.strip() for line in f if line.strip()}
            except Exception:
                outdated_set = set()

        import json as _json
        all_entries = []
        for json_file in sorted(confluence_docs_dir.glob("*.json")):
            try:
                with open(json_file, "r", encoding="utf-8") as f:
                    docs = _json.load(f)
            except Exception:
                continue
            if not isinstance(docs, list):
                docs = [docs]
            for d in docs:
                src = d.get("source") or {}
                pid = src.get("page_id", "") or ""
                all_entries.append(
                    {
                        "Файл": json_file.name,
                        "Заголовок": d.get("title", "—"),
                        "Ссылка": src.get("url", ""),
                        "page_id": pid,
                        "outdated": pid in outdated_set,
                    }
                )
        if not all_entries:
            st.info("В `confluence_docs/` пока нет JSON документов.")
            return

        total = len(all_entries)
        in_use = sum(1 for e in all_entries if not e["outdated"])
        st.metric("Документов в KB (Confluence/manual)", total)
        st.caption(f"Актуальных (попадут в поиск): {in_use}, устаревших: {total - in_use}")

        with st.expander("Показать список документов: обновить / пометить устаревшим"):
            for i, row in enumerate(all_entries, 1):
                label = f"**{i}. {row['Заголовок']}**"
                if row["outdated"]:
                    label += " — ⚠️ устаревший"
                st.markdown(label)
                if row["Ссылка"]:
                    st.caption(f"📎 [{row['Ссылка']}]({row['Ссылка']})")
                else:
                    st.caption(f"Файл: {row['Файл']}, page_id: `{row['page_id']}`")
                if not row["page_id"]:
                    continue
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("🔄 Обновить из Confluence", key=f"upd_{i}_{row['page_id']}"):
                        if HAS_CONFLUENCE_GENERATOR:
                            with st.spinner("Обновление..."):
                                try:
                                    gen = ConfluenceKBGenerator(client=_get_client())
                                    n = gen.update_docs_by_page_ids([row["page_id"]])
                                    st.success(f"Обновлено документов: {n}. Нажмите «Перезагрузить KB в Qdrant» для применения.")
                                except Exception as e:
                                    st.error(str(e))
                            st.rerun()
                        else:
                            st.error("Модуль синхронизации недоступен.")
                with col2:
                    if not row["outdated"] and st.button("📌 Пометить устаревшим", key=f"out_{i}_{row['page_id']}"):
                        try:
                            gen = ConfluenceKBGenerator(client=_get_client()) if HAS_CONFLUENCE_GENERATOR else None
                            if gen:
                                gen.add_to_outdated(row["page_id"])
                                st.success("Помечен устаревшим. Перезагрузите KB в Qdrant, чтобы исключить из поиска.")
                            else:
                                path = confluence_docs_dir / "outdated.txt"
                                with open(path, "a", encoding="utf-8") as f:
                                    f.write(row["page_id"] + "\n")
                                st.success("Помечен устаревшим.")
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))
                with col3:
                    if row["outdated"] and st.button("✅ Вернуть в актуальные", key=f"rev_{i}_{row['page_id']}"):
                        try:
                            if HAS_CONFLUENCE_GENERATOR:
                                gen = ConfluenceKBGenerator(client=_get_client())
                                gen.remove_from_outdated(row["page_id"])
                            else:
                                path = confluence_docs_dir / "outdated.txt"
                                if path.exists():
                                    lines = [l for l in open(path, encoding="utf-8") if l.strip() != row["page_id"]]
                                    with open(path, "w", encoding="utf-8") as f:
                                        f.writelines(lines)
                                st.success("Снята пометка. Перезагрузите KB в Qdrant.")
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))

    def _render_confluence_sync_and_reload():
        st.subheader("Confluence")
        # Настройки (можно переопределить из config.env)
        confluence_url = st.text_input(
            "URL Confluence",
            value=os.getenv("CONFLUENCE_URL", "https://docs.steccom.ru"),
            key="confluence_url_lib",
            help="Базовый URL Confluence (например https://docs.steccom.ru)",
        )
        confluence_token = st.text_input(
            "Токен (Personal Access Token)",
            value=os.getenv("CONFLUENCE_TOKEN", ""),
            type="password",
            key="confluence_token_lib",
            help="Задаётся в config.env или здесь. Не сохраняйте в коде.",
        )
        if not confluence_token:
            st.caption("💡 Задайте CONFLUENCE_TOKEN в config.env или введите выше.")

        client = _get_client(confluence_url, confluence_token)

        if st.button("🔌 Проверить подключение к Confluence", type="primary", key="confluence_check_btn"):
            with st.spinner("Проверка..."):
                ok, msg = client.check_connection()
            if ok:
                st.success(msg)
                try:
                    spaces = client.get_spaces(limit=20)
                    if spaces:
                        st.subheader("Пространства (первые 20)")
                        for s in spaces:
                            st.text(f"  • {s.get('key', '')} — {s.get('name', '')}")
                except Exception as e:
                    st.warning(f"Список пространств: {e}")
            else:
                st.error(msg)

        st.markdown("---")
        st.subheader("Синхронизация пространства в KB")
        st.caption(
            "Ключ пространства — это не URL. Для пользователя из ссылки вида "
            "«.../username=n.shiriaev» укажите ключ личного пространства: **~n.shiriaev** (тильда + логин)."
        )
        space_key = st.text_input(
            "Ключ пространства (Space key)",
            value="",
            placeholder="например DEMO, SPC или ~n.shiriaev",
            key="confluence_space_key",
        )
        limit_pages = st.number_input(
            "Макс. страниц за один запуск (0 = без ограничения)",
            min_value=0,
            value=50,
            key="confluence_limit",
        )
        limit = None if limit_pages == 0 else int(limit_pages)

        if st.button("📥 Синхронизировать пространство в KB", key="confluence_sync_btn"):
            if not HAS_CONFLUENCE_GENERATOR:
                st.error("Модуль синхронизации недоступен. Разверните `kb_billing/rag/confluence_kb_generator.py` и установите beautifulsoup4.")
            elif not space_key.strip():
                st.error("Введите ключ пространства.")
            else:
                with st.spinner("Синхронизация..."):
                    try:
                        gen = ConfluenceKBGenerator(client=client)
                        docs = gen.sync_space(space_key.strip(), limit=limit)
                        out_dir = gen.get_synced_docs_path()
                        st.success(f"Сохранено документов: **{len(docs)}** в `{out_dir}`")
                        if docs:
                            with st.expander("Первые заголовки"):
                                for d in docs[:15]:
                                    st.text(f"  • {d.get('title', '')}")
                    except Exception as e:
                        err = str(e)
                        st.error(err)
                        if "404" in err and "~" in (space_key or ""):
                            st.info(
                                "Для личных пространств этот Confluence может не отдавать страницы по spaceKey. "
                                "Используйте блок ниже «Конкретные страницы по URL или ID»: вставьте ссылки на нужные страницы."
                            )
                        import traceback
                        st.code(traceback.format_exc())

        st.markdown("---")
        st.subheader("Конкретные страницы по URL или ID")
        st.caption("По одному URL или Page ID на строку. К каждой строке можно добавить описание для удобства (таб или два пробела): «URL  Краткое описание для ссылки инженера» — описание сохранится в документе.")
        try:
            from kb_billing.rag.vision_annotation import is_vision_available
            if is_vision_available():
                st.info("✅ **Аннотация изображений включена:** при синхронизации картинки и схемы будут описаны через Gemini Vision (ключ задан в config.env).")
            else:
                st.warning(
                    "⚠️ **Аннотация изображений недоступна.** Картинки/схемы сохранятся без текстового описания. "
                    "Чтобы включить: в **config.env** (в каталоге приложения, на сервере — где запущен Streamlit) задайте **GEMINI_API_KEY** или **OPENAI_API_KEY** (ключ ProxyAPI.ru). "
                    "После сохранения config.env перезапустите приложение и снова выполните синхронизацию страницы."
                )
        except Exception:
            st.warning("⚠️ Проверка аннотации изображений недоступна (модуль vision_annotation не загружен).")
        page_urls_or_ids = st.text_area(
            "URL страниц/вложений или Page ID (каждый с новой строки)",
            value="",
            height=120,
            placeholder="https://docs.steccom.ru/download/attachments/4392673/схема.png  Схема сети VSAT\n123457  Инструкция iDirect",
            key="confluence_page_urls",
        )
        merge_pages = st.checkbox(
            "Дополнить/обновить существующий список (merge)",
            value=False,
            key="confluence_merge_pages",
            help="Если включено, новые страницы добавятся к уже сохранённым в custom_pages; совпадающие по page_id обновятся.",
        )
        with_children = st.checkbox(
            "Загрузить также дочерние страницы (рекурсивно по API)",
            value=False,
            key="confluence_with_children",
            help="Для каждой указанной страницы подтянуть все прямые дочерние (например, страница «Спецификации» и все спецификации под ней).",
        )
        if st.button("📥 Синхронизировать выбранные страницы в KB", key="confluence_sync_pages_btn"):
            if not HAS_CONFLUENCE_GENERATOR:
                st.error("Модуль синхронизации недоступен.")
            elif not page_urls_or_ids.strip():
                st.error("Введите хотя бы один URL или ID страницы.")
            else:
                lines = [s.strip() for s in page_urls_or_ids.strip().splitlines() if s.strip()]
                with st.spinner("Синхронизация страниц..."):
                    try:
                        from kb_billing.rag.confluence_sync_runner import sync_pages
                        out_path, count = sync_pages(
                            lines,
                            output_dir=confluence_docs_dir,
                            output_suffix="custom_pages",
                            merge=merge_pages,
                            client=client,
                            include_children=with_children,
                            max_depth=1,
                            limit_children=200,
                        )
                        st.success(f"Сохранено документов: **{count}** в `{out_path.name}`")
                        if count > 0 and out_path.exists():
                            import json as _json
                            try:
                                with open(out_path, "r", encoding="utf-8") as f:
                                    docs = _json.load(f)
                                if isinstance(docs, list) and docs:
                                    with st.expander("Заголовки"):
                                        for d in docs[:20]:
                                            comment = (d.get("source") or {}).get("engineer_comment", "")
                                            if comment:
                                                st.text(f"  • {d.get('title', '')}  — {comment}")
                                            else:
                                                st.text(f"  • {d.get('title', '')}")
                            except Exception:
                                pass
                    except Exception as e:
                        st.error(str(e))
                        import traceback
                        st.code(traceback.format_exc())

        st.markdown("---")
        st.subheader("Обновление спутниковой части KB в Qdrant")
        st.markdown(
            "После синхронизации из Confluence или правок в «Редактор KB» обновите **только документы Confluence** в Qdrant. "
            "Данные биллинга (Q/A примеры, таблицы, представления) не трогаются — их перезагружают во вкладке **«KB Библиотекарь»**."
        )
        _json_files = list(confluence_docs_dir.glob("*.json")) if confluence_docs_dir.exists() else []
        try:
            _coll = os.getenv("QDRANT_COLLECTION", "kb_billing")
        except Exception:
            _coll = "kb_billing"
        st.caption(
            f"Каталог загрузки: **{confluence_docs_dir.resolve()}** · JSON: **{len(_json_files)}** "
            + (f"({', '.join(f.name for f in _json_files[:5])}" + (" …" if len(_json_files) > 5 else "") + ")" if _json_files else "(пусто)")
            + f" · Коллекция Qdrant: **{_coll}**"
        )
        if st.button("🔄 Перезагрузить в Qdrant только Confluence (спутниковая KB)", key="confluence_reload_kb_btn"):
            with st.spinner("Загрузка в Qdrant (только Confluence)..."):
                try:
                    from kb_billing.rag.kb_loader import KBLoader
                    # Тот же kb_dir, что и для списка документов (confluence_docs_dir), чтобы не разъезжаться при разделении доменов
                    loader = KBLoader(kb_dir=confluence_docs_dir.parent)
                    n = loader.reload_confluence_only()
                    st.success(f"В Qdrant обновлено **{n}** чанков Confluence. Данные биллинга не изменялись.")
                    if n == 0 and _json_files:
                        st.warning(
                            "Загружено 0 чанков при наличии JSON. Проверьте: 1) все page_id не в outdated.txt? "
                            "2) в каждом документе есть секции с непустым текстом."
                        )
                except Exception as e:
                    st.error(str(e))
                    import traceback
                    st.code(traceback.format_exc())

    def _render_answer_testing():
        st.subheader("💬 Тест ответов спутникового ассистента")
        st.caption("Через ProxyAPI.ru: задайте GEMINI_API_KEY или OPENAI_API_KEY в config.env (по умолчанию https://api.proxyapi.ru/google).")
        if sys.version_info < (3, 9):
            st.warning(
                "Для Спутникового ассистента (Gemini) нужен **Python 3.9+**. Сейчас: **%s**. "
                "На сервере запустите приложение под Python 3.9+ (например: `conda activate py11` или создайте venv с python3.9)."
                % (sys.version.split()[0],)
            )
            return
        try:
            from kb_billing.rag.satellite_librarian_agent import SatelliteLibrarianAgent, ENGINEER_PROMPT, SUBSCRIBER_PROMPT
            from kb_billing.rag.rag_assistant import RAGAssistant
            has_agent = True
        except Exception:
            has_agent = False

        if not has_agent:
            st.error(
                "Не удалось загрузить модули ассистента. Установите зависимости (в т.ч. **httpx**). "
                "В config.env задайте **GEMINI_API_KEY** или **OPENAI_API_KEY** (ключ ProxyAPI.ru). "
                "По умолчанию используется https://api.proxyapi.ru/google (как в brats)."
            )
            return

        persona = st.radio("Режим", ["Инженер", "Абонент"], horizontal=True, key="sat_persona")
        key_prefix = "sat_eng" if persona == "Инженер" else "sat_sub"
        sys_prompt = ENGINEER_PROMPT if persona == "Инженер" else SUBSCRIBER_PROMPT
        avatar_user = "👷" if persona == "Инженер" else "🙂"
        avatar_assistant = "🛰️"

        chat_key = f"{key_prefix}_chat"
        input_key = f"{key_prefix}_input"
        chunks_key = f"{key_prefix}_last_chunks"
        use_rag_key = f"{key_prefix}_use_rag"

        # История диалога: при первом заходе подгружаем из БД (если есть пользователь и модуль)
        if chat_key not in st.session_state:
            if HAS_CHAT_DB and satellite_chat_db:
                username = st.session_state.get("username") or ""
                if username:
                    try:
                        st.session_state[chat_key] = satellite_chat_db.get_history(username, key_prefix)
                    except Exception:
                        st.session_state[chat_key] = []
                else:
                    st.session_state[chat_key] = []
            else:
                st.session_state[chat_key] = []
        if chunks_key not in st.session_state:
            st.session_state[chunks_key] = None

        # Применить отложенное значение поля ввода до отрисовки виджета (избегаем "cannot be modified after widget")
        input_pending_key = f"{key_prefix}_input_pending"
        last_error_key = f"{key_prefix}_last_error"
        if input_pending_key in st.session_state:
            st.session_state[input_key] = st.session_state[input_pending_key]
            del st.session_state[input_pending_key]
        if input_key not in st.session_state:
            st.session_state[input_key] = ""

        # Показать ошибку после rerun (если была при прошлой отправке)
        if last_error_key in st.session_state:
            err_msg, err_tb = st.session_state[last_error_key]
            del st.session_state[last_error_key]
            st.error(err_msg)
            st.code(err_tb)

        use_rag = st.checkbox("Подмешивать в контекст фрагменты из KB (RAG)", value=True, key=use_rag_key)
        if HAS_CHAT_DB and st.session_state.get("username"):
            st.caption("История диалога сохраняется в БД — контекст не теряется при перезагрузке страницы.")

        for idx, msg in enumerate(st.session_state[chat_key]):
            if msg["role"] == "user":
                with st.chat_message("user", avatar=avatar_user):
                    st.markdown(msg["content"])
            else:
                with st.chat_message("assistant", avatar=avatar_assistant):
                    st.markdown(msg["content"])
                    if idx == len(st.session_state[chat_key]) - 1 and st.session_state[chunks_key]:
                        with st.expander("📋 Какие чанки использованы для этого ответа"):
                            for i, ch in enumerate(st.session_state[chunks_key], 1):
                                st.markdown(
                                    f"**{i}. Документ:** {ch.get('title', '—')} · **Секция:** {ch.get('section_title', '—')} · **Релевантность:** {ch.get('similarity', 0):.2%}"
                                )
                                if ch.get("source_url"):
                                    st.caption(f"[Открыть в Confluence]({ch['source_url']})")
                                preview = (ch.get("content") or "")
                                st.text(preview[:300] + ("..." if len(preview) > 300 else ""))

        # Спецификация из последнего ответа: таблица на экране + скачивание в XLSX
        last_assistant_content = next(
            (m["content"] for m in reversed(st.session_state[chat_key]) if m["role"] == "assistant"),
            None,
        )
        if last_assistant_content and len(st.session_state[chat_key]) > 0:
            spec_title, spec_rows = _parse_spec_table(last_assistant_content)
            if spec_rows:
                try:
                    import pandas as pd
                    df_spec = pd.DataFrame(spec_rows, columns=["Наименование", "Количество"])
                    st.subheader("📋 Спецификация (таблица)")
                    if spec_title:
                        st.caption(spec_title)
                    st.dataframe(df_spec, use_container_width=True, hide_index=True)
                    col_xlsx, col_docx = st.columns(2)
                    with col_xlsx:
                        buf = io.BytesIO()
                        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                            df_spec.to_excel(writer, index=False, sheet_name="Спецификация")
                        buf.seek(0)
                        st.download_button(
                            "⬇ Скачать как .xlsx",
                            data=buf,
                            file_name="спека_судна.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"{key_prefix}_download_spec_xlsx",
                        )
                    with col_docx:
                        try:
                            docx_bytes = _build_spec_docx(spec_title or "", spec_rows)
                            st.download_button(
                                "⬇ Скачать как .docx (Word)",
                                data=docx_bytes,
                                file_name="спека_судна.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"{key_prefix}_download_spec_docx",
                            )
                        except Exception:
                            pass  # Word опционален
                except Exception:
                    st.download_button(
                        "⬇ Скачать последний ответ как .txt",
                        data=last_assistant_content,
                        file_name="спека_судна.txt",
                        mime="text/plain; charset=utf-8",
                        key=f"{key_prefix}_download_last_reply",
                    )
            else:
                st.download_button(
                    "⬇ Скачать последний ответ как .txt",
                    data=last_assistant_content,
                    file_name="спека_судна.txt",
                    mime="text/plain; charset=utf-8",
                    key=f"{key_prefix}_download_last_reply",
                )

        user_text = st.text_area(
            "Ваш вопрос",
            height=120,
            placeholder="Опишите проблему / задайте вопрос...",
            key=input_key,
        )
        col_send, col_clear, col_drawio, col_drawio_kb = st.columns([1, 1, 1, 1])
        with col_send:
            send_clicked = st.button("📤 Отправить", type="primary", key=f"{key_prefix}_send_btn")
        with col_clear:
            clear_clicked = st.button("🗑 Очистить диалог", key=f"{key_prefix}_clear_btn")
        drawio_key = f"{key_prefix}_drawio_xml"
        with col_drawio:
            gen_drawio_clicked = st.button(
                "📐 Создать draw.io по диалогу",
                key=f"{key_prefix}_gen_drawio_btn",
                disabled=len(st.session_state[chat_key]) == 0,
                help="После уточняющих вопросов сгенерировать схему draw.io по спецификации из диалога",
            )
        last_user_msg = next((m["content"] for m in reversed(st.session_state[chat_key]) if m["role"] == "user"), "")
        url_match = re.search(r"/download/attachments/(\d+)/([^/?]+)", last_user_msg, re.IGNORECASE) if last_user_msg else None
        draw_scheme_words = "нарисуй" in last_user_msg.lower() or "нарисовать" in last_user_msg.lower()
        with col_drawio_kb:
            gen_drawio_kb_clicked = st.button(
                "📐 Draw.io из KB по ссылке",
                key=f"{key_prefix}_gen_drawio_kb_btn",
                disabled=not (url_match and draw_scheme_words and len(st.session_state[chat_key]) > 0),
                help="Сгенерировать draw.io по описанию схемы из KB (аннотация Gemini по этой картинке)",
            )

        if gen_drawio_kb_clicked and url_match and st.session_state[chat_key]:
            page_id, filename = url_match.group(1), url_match.group(2).strip()
            with st.spinner("Загрузка описания из KB и генерация draw.io..."):
                try:
                    rag = RAGAssistant()
                    chunks = rag.get_confluence_chunks_by_page_id(page_id, limit=50)
                    spec_chunk = None
                    for ch in chunks:
                        stitle = (ch.get("section_title") or "")
                        if filename in stitle and (stitle.startswith("Вложение:") or "Вложение" in stitle):
                            content = (ch.get("content") or "").strip()
                            if len(content) > 150 and "OCR не извлечён" not in content and "vision API" not in content:
                                spec_chunk = content
                                break
                    if not spec_chunk:
                        st.error("В KB нет текстового описания этой картинки (аннотация Vision). Синхронизируйте страницу с включённой аннотацией и перезагрузите Confluence в Qdrant.")
                    else:
                        agent = SatelliteLibrarianAgent(rag_assistant=rag, system_prompt=sys_prompt)
                        history = [(m["role"], m["content"]) for m in st.session_state[chat_key]][-10:]
                        spec = f"Спецификация схемы из базы знаний (аннотация изображения {filename}):\n\n{spec_chunk}"
                        xml_content = agent.generate_drawio(spec, history=history)
                        st.session_state[drawio_key] = xml_content
                        st.success("Draw.io сгенерирован по описанию из KB. Скачайте ниже.")
                        st.rerun()
                except Exception as e:
                    import traceback
                    st.error("Ошибка: " + str(e))
                    st.code(traceback.format_exc())

        if gen_drawio_clicked and st.session_state[chat_key]:
            use_rag = st.session_state.get(use_rag_key, True)
            with st.spinner("Генерация draw.io по диалогу..."):
                try:
                    rag = RAGAssistant() if use_rag else None
                    agent = SatelliteLibrarianAgent(rag_assistant=rag, system_prompt=sys_prompt)
                    chat_list = st.session_state[chat_key]
                    last_user = next((m["content"] for m in reversed(chat_list) if m["role"] == "user"), "Схема сети по диалогу")
                    spec = "Итоговое пожелание пользователя: " + last_user
                    history = [(m["role"], m["content"]) for m in chat_list]
                    xml_content = agent.generate_drawio(spec, history=history)
                    st.session_state[drawio_key] = xml_content
                    st.rerun()
                except Exception as e:
                    import traceback
                    st.error("Ошибка генерации draw.io: " + str(e))
                    st.code(traceback.format_exc())

        if clear_clicked and st.session_state[chat_key]:
            if HAS_CHAT_DB and satellite_chat_db:
                username = st.session_state.get("username") or ""
                if username:
                    try:
                        satellite_chat_db.clear_history(username, key_prefix)
                    except Exception:
                        pass
            st.session_state[chat_key] = []
            st.session_state[chunks_key] = None
            if drawio_key in st.session_state:
                del st.session_state[drawio_key]
            st.rerun()

        if drawio_key in st.session_state and st.session_state[drawio_key]:
            with st.expander("📐 Сгенерированный draw.io (скачать или скопировать XML)", expanded=True):
                xml_content = st.session_state[drawio_key]
                if xml_content.strip().startswith("<"):
                    st.download_button(
                        "⬇ Скачать как .drawio",
                        data=xml_content,
                        file_name="scheme.drawio",
                        mime="application/xml",
                        key=f"{key_prefix}_download_drawio",
                    )
                    st.code(xml_content[:8000] + ("\n... (обрезано)" if len(xml_content) > 8000 else ""), language="xml")
                else:
                    st.warning("Ответ не похож на XML (возможно, ошибка модели):")
                    st.text(xml_content[:1500])

        if send_clicked and (user_text or "").strip():
            q = (user_text or "").strip()
            st.session_state[chat_key].append({"role": "user", "content": q})
            st.session_state[chunks_key] = None
            input_pending_key = f"{key_prefix}_input_pending"
            last_error_key = f"{key_prefix}_last_error"
            with st.spinner("Ассистент думает..."):
                try:
                    rag = RAGAssistant() if use_rag else None
                    chunks = None
                    if use_rag and rag:
                        chunks = rag.search_semantic(q, content_type="confluence_section", limit=10)
                        if not chunks:
                            chunks = rag.search_semantic(q, content_type="confluence_doc", limit=10)
                        st.session_state[chunks_key] = chunks
                    agent = SatelliteLibrarianAgent(rag_assistant=rag, system_prompt=sys_prompt)
                    history = [(m["role"], m["content"]) for m in st.session_state[chat_key][:-1]][-20:]
                    answer = agent.ask(q, use_rag=use_rag, history=history if history else None)
                    st.session_state[chat_key].append({"role": "assistant", "content": answer})
                    # Сохранить реплики в БД для сохранения контекста между сеансами
                    if HAS_CHAT_DB and satellite_chat_db:
                        username = st.session_state.get("username") or ""
                        if username:
                            try:
                                satellite_chat_db.append_message(username, key_prefix, "user", q)
                                satellite_chat_db.append_message(username, key_prefix, "assistant", answer)
                            except Exception:
                                pass
                    st.session_state[input_pending_key] = ""  # очистить поле на следующем run (нельзя менять input_key после виджета)
                    st.rerun()
                except Exception as e:
                    import traceback
                    st.session_state[chat_key].pop()
                    st.session_state[input_pending_key] = q  # вернуть текст при ошибке (применится до виджета при rerun)
                    st.session_state[last_error_key] = (str(e), traceback.format_exc())
                    st.rerun()

    with tab_chat:
        _render_answer_testing()
    with tab_retrieval:
        _render_retrieval_test()
    with tab_editor:
        _render_manual_kb_editor()
        st.markdown("---")
        _render_confluence_docs_list()
    with tab_sync:
        _render_confluence_sync_and_reload()
