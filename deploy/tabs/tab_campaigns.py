"""
Закладка: Рекламные email кампании
"""
import streamlit as st
from pathlib import Path
from datetime import datetime
import io
import re
import smtplib
import time
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import cx_Oracle
from typing import List, Optional, Tuple

# Импортируем функции безопасной рассылки
try:
    import sys
    from pathlib import Path as PathLib
    # Добавляем путь к корню проекта для импорта common
    project_root = PathLib(__file__).parent.parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))
    from tabs.common import send_email_safely
except ImportError:
    # Если не удалось импортировать, будем использовать старую логику
    send_email_safely = None


def _get_db_connection():
    """Получить подключение к Oracle"""
    from utils.db_connection import get_db_connection
    return get_db_connection()


def _parse_email_list(email_text) -> List[str]:
    """
    Парсинг списка email из текста (разделитель - запятая, точка с запятой или новая строка)
    Возвращает список валидных email адресов
    
    Поддерживает строки и cx_Oracle.LOB объекты
    """
    # Если это LOB объект, читаем его в строку
    if email_text is None:
        return []
    
    # Проверяем, является ли это LOB объектом (cx_Oracle.LOB)
    if hasattr(email_text, 'read'):
        try:
            email_text = email_text.read()
        except Exception:
            email_text = str(email_text)
    
    # Конвертируем в строку, если еще не строка
    if not isinstance(email_text, str):
        email_text = str(email_text) if email_text else ''
    
    # Убираем пробелы по краям и проверяем на пустую строку или только пробел
    email_text = email_text.strip()
    if not email_text or email_text == '':
        return []
    
    # Заменяем переносы строк и точку с запятой на запятые
    email_text = email_text.replace('\n', ',').replace(';', ',')
    # Разбиваем по запятым
    emails = [e.strip() for e in email_text.split(',')]
    # Фильтруем пустые и валидируем email
    valid_emails = []
    email_pattern = re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')
    for email in emails:
        email = email.strip()
        if email and email_pattern.match(email):
            valid_emails.append(email.lower())
    return valid_emails


def _looks_like_html(text: str) -> bool:
    """Проверяет, содержит ли строка HTML-теги (для выбора режима: HTML как есть или экранирование)."""
    if not text or not text.strip():
        return False
    # Проверка на типичные теги: открывающие/закрывающие или самозакрывающиеся
    tag_pattern = re.compile(
        r'<(?:p|div|span|b|strong|i|em|u|br|font|h[1-6]|ul|ol|li|table|tr|td|th|a)\s*(?:\s[^>]*)?/?>',
        re.IGNORECASE
    )
    return bool(tag_pattern.search(text))


def _remove_duplicate_text(text: str) -> str:
    """
    Удалить дублирование текста из строки.
    Проверяет несколько методов для обнаружения дублирования.
    """
    if not text or len(text) < 10:
        return text
    
    text = text.strip()
    original_length = len(text)
    
    if original_length < 20:
        return text
    
    # Метод 1: Ищем повторение по маркеру "Уважаемые коллеги"
    marker_phrase = "Уважаемые коллеги"
    if marker_phrase in text:
        positions = []
        start = 0
        while True:
            pos = text.find(marker_phrase, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + 1
        
        if len(positions) >= 2:
            first_occurrence = positions[0]
            second_occurrence = positions[1]
            first_part = text[:second_occurrence].strip()
            second_part = text[second_occurrence:].strip()
            
            marker_len = len(marker_phrase)
            first_after_marker = first_part[first_occurrence + marker_len:].strip()
            second_after_marker = second_part[marker_len:].strip() if len(second_part) > marker_len else second_part.strip()
            
            if len(first_after_marker) > 50 and len(second_after_marker) > 50:
                min_len = min(len(first_after_marker), len(second_after_marker))
                matches = sum(1 for a, b in zip(first_after_marker[:min_len], second_after_marker[:min_len]) if a == b)
                if matches > min_len * 0.9:  # 90% совпадение
                    text = first_part
                    logging.info(f"Обнаружено дублирование по маркеру '{marker_phrase}' (было {original_length} -> стало {len(text)}), исправлено")
                    return text
    
    # Метод 2: Проверяем точное дублирование пополам
    half_length = original_length // 2
    first_half = text[:half_length].strip()
    second_half = text[half_length:].strip()
    
    if first_half == second_half and len(first_half) > 10:
        text = first_half
        logging.info(f"Обнаружено точное дублирование пополам (было {original_length} -> стало {len(text)}), исправлено")
        return text
    
    # Метод 3: Проверяем дублирование по предложениям (разделитель - двойной перенос строки)
    sentences = text.split('\n\n')
    if len(sentences) >= 4:
        mid_point = len(sentences) // 2
        first_part = '\n\n'.join(sentences[:mid_point]).strip()
        second_part = '\n\n'.join(sentences[mid_point:]).strip()
        if first_part == second_part and len(first_part) > 20:
            text = first_part
            logging.info(f"Обнаружено дублирование по предложениям (было {original_length} -> стало {len(text)}), исправлено")
            return text
    
    return text


def _extract_subject_and_greeting_from_docx(docx_content: bytes) -> Tuple[str, str]:
    """
    Извлечь тему письма (subject) и приветствие (greeting) из DOCX файла.
    Предполагается, что первая строка - это тема, остальное - приветствие.
    Автоматически удаляет дублирование текста из greeting.
    Возвращает (subject, greeting)
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_content))
        
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        if not paragraphs:
            return "", ""
        
        # Первая непустая строка - тема письма
        subject = paragraphs[0] if paragraphs else ""
        
        # Остальные строки - приветствие
        greeting_lines = paragraphs[1:] if len(paragraphs) > 1 else []
        greeting = '\n\n'.join(greeting_lines) if greeting_lines else ""
        
        # Удаляем дублирование из greeting перед возвратом
        if greeting:
            greeting = _remove_duplicate_text(greeting)
        
        return subject, greeting
    except ImportError:
        st.error("⚠️ Библиотека python-docx не установлена. Установите: pip install python-docx")
        return "", ""
    except Exception as e:
        st.warning(f"⚠️ Ошибка при извлечении текста из DOCX: {e}")
        return "", ""


def _docx_to_html(docx_content: bytes) -> str:
    """
    Конвертация DOCX в HTML для отправки по email
    Использует python-docx для чтения документа
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        doc = Document(io.BytesIO(docx_content))
        html_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Простое форматирование: жирный текст
                text = paragraph.text
                # Заменяем переносы строк на <br>
                text = text.replace('\n', '<br>')
                html_parts.append(f"<p>{text}</p>")
        
        # Обработка таблиц если есть
        for table in doc.tables:
            html_parts.append("<table border='1' style='border-collapse: collapse; margin: 10px 0;'>")
            for row in table.rows:
                html_parts.append("<tr>")
                for cell in row.cells:
                    html_parts.append(f"<td style='padding: 5px;'>{cell.text}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")
        
        return '\n'.join(html_parts)
    except ImportError:
        st.error("⚠️ Библиотека python-docx не установлена. Установите: pip install python-docx")
        return "<p>Ошибка: библиотека python-docx не установлена</p>"
    except Exception as e:
        st.warning(f"⚠️ Ошибка при конвертации DOCX: {e}. Будет отправлен пустой текст.")
        return "<p>Ошибка при обработке документа</p>"


def _docx_to_text(docx_content: bytes) -> str:
    """Конвертация DOCX в простой текст (fallback)"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_content))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        return '\n\n'.join(text_parts)
    except Exception as e:
        return f"Ошибка при обработке документа: {e}"


def _save_campaign_to_db(
    conn,
    campaign_name: str,
    subject: str,
    greeting: str,
    email_list: List[str],
    docx_content: Optional[bytes],
    docx_filename: Optional[str],
    created_by: str,
    test_mode: bool = False,
    test_emails: Optional[List[str]] = None
) -> Optional[int]:
    """
    Сохранить кампанию в базу данных Oracle
    Возвращает CAMPAIGN_ID или None при ошибке
    """
    try:
        # Убираем дублирование только для обычного текста (не для HTML, чтобы не повредить разметку)
        if greeting and not _looks_like_html(greeting):
            greeting = _remove_duplicate_text(greeting)
        
        cursor = conn.cursor()
        # Преобразуем список email в строку
        # Для тестовой рассылки основной список может быть пустым, но поле не может быть NULL
        if email_list and len(email_list) > 0:
            email_list_str = ','.join(email_list)
        else:
            # Если список пуст, используем пробел вместо пустой строки
            # Oracle может интерпретировать пустую строку '' как NULL для CLOB,
            # поэтому используем пробел как минимальное непустое значение
            email_list_str = ' '
        
        test_emails_str = ','.join(test_emails) if test_emails else None
        
        # Создаем переменную для получения ID
        campaign_id_var = cursor.var(cx_Oracle.NUMBER)
        
        if docx_content:
            cursor.execute("""
                INSERT INTO EMAIL_CAMPAIGNS (
                    CAMPAIGN_NAME, SUBJECT, GREETING, EMAIL_LIST,
                    DOCX_CONTENT, DOCX_FILENAME, EMAILS_TOTAL,
                    CREATED_BY, STATUS, TEST_MODE, TEST_EMAILS
                ) VALUES (
                    :1, :2, :3, :4, :5, :6, :7, :8, 'DRAFT', :9, :10
                )
                RETURNING CAMPAIGN_ID INTO :11
            """, (
                campaign_name, subject, greeting, email_list_str,
                cx_Oracle.Binary(docx_content), docx_filename,
                len(email_list), created_by, 1 if test_mode else 0, test_emails_str,
                campaign_id_var
            ))
        else:
            cursor.execute("""
                INSERT INTO EMAIL_CAMPAIGNS (
                    CAMPAIGN_NAME, SUBJECT, GREETING, EMAIL_LIST,
                    EMAILS_TOTAL, CREATED_BY, STATUS, TEST_MODE, TEST_EMAILS
                ) VALUES (
                    :1, :2, :3, :4, :5, :6, 'DRAFT', :7, :8
                )
                RETURNING CAMPAIGN_ID INTO :9
            """, (
                campaign_name, subject, greeting, email_list_str,
                len(email_list), created_by, 1 if test_mode else 0, test_emails_str,
                campaign_id_var
            ))
        
        campaign_id = campaign_id_var.getvalue()[0]
        conn.commit()
        cursor.close()
        return campaign_id
    except Exception as e:
        st.error(f"❌ Ошибка при сохранении кампании в БД: {e}")
        conn.rollback()
        return None


def _send_email_campaign(
    conn,
    campaign_id: int,
    smtp_host: str = 'mail.steccom.ru',
    smtp_port: int = 25,
    from_email: str = 'sales@steccom.ru',
    smtp_password: Optional[str] = None,
    test_mode: bool = False,
    test_emails: Optional[List[str]] = None,
    sent_by: Optional[str] = None,
    delay_between_emails: float = 2.0,
    delay_after_batch: float = 60.0,
    batch_size: int = 10
) -> Tuple[int, int, str]:
    """
    Отправить кампанию по email
    Возвращает (отправлено, ошибок, сообщение)
    """
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT CAMPAIGN_NAME, SUBJECT, GREETING, EMAIL_LIST,
                   DOCX_CONTENT, DOCX_FILENAME, EMAILS_TOTAL, TEST_MODE, TEST_EMAILS
            FROM EMAIL_CAMPAIGNS
            WHERE CAMPAIGN_ID = :1
        """, (campaign_id,))
        
        row = cursor.fetchone()
        if not row:
            return 0, 0, "Кампания не найдена"
        
        campaign_name, subject, greeting, email_list_str, docx_content, docx_filename, emails_total, db_test_mode, db_test_emails = row
        
        # Конвертируем ВСЕ LOB объекты в строки ДО использования len() или других операций
        # Функция для безопасной конвертации LOB в строку
        def convert_lob_to_string(lob_value, default=''):
            if lob_value is None:
                return default
            if isinstance(lob_value, cx_Oracle.LOB):
                try:
                    return lob_value.read()
                except Exception as e:
                    logging.warning(f"Ошибка при чтении LOB: {e}")
                    return str(lob_value) if lob_value else default
            elif hasattr(lob_value, 'read') and hasattr(lob_value, 'size'):
                try:
                    return lob_value.read()
                except Exception as e:
                    logging.warning(f"Ошибка при чтении LOB-подобного объекта: {e}")
                    return str(lob_value) if lob_value else default
            elif not isinstance(lob_value, str):
                return str(lob_value) if lob_value else default
            return lob_value
        
        # Конвертируем все поля, которые могут быть LOB
        email_list_str = convert_lob_to_string(email_list_str, '')
        greeting = convert_lob_to_string(greeting, '')
        db_test_emails = convert_lob_to_string(db_test_emails, None) if db_test_emails is not None else None
        
        # Определяем режим рассылки: если передан test_mode или в БД установлен TEST_MODE=1
        use_test_mode = test_mode or (db_test_mode == 1)
        
        # Выбираем список email для рассылки
        if use_test_mode:
            # Используем тестовые email из параметра или из БД
            if test_emails and len(test_emails) > 0:
                # Если передан список тестовых email, используем его
                email_list = test_emails
            elif db_test_emails:
                # Если в БД есть тестовые email, парсим их
                parsed_test_emails = _parse_email_list(db_test_emails)
                if parsed_test_emails and len(parsed_test_emails) > 0:
                    email_list = parsed_test_emails
                else:
                    return 0, 0, "Для тестовой рассылки не указаны контрольные email (список пуст или невалиден)"
            else:
                return 0, 0, "Для тестовой рассылки не указаны контрольные email"
        else:
            # Обычная рассылка по основному списку
            email_list = _parse_email_list(email_list_str)
        
        # Финальная проверка списка
        if not email_list or len(email_list) == 0:
            mode_text = "тестовой" if use_test_mode else "обычной"
            return 0, 0, f"Список email для {mode_text} рассылки пуст или невалиден"
        
        # Читаем BLOB вложения (PDF) из Oracle, если есть
        attachment_content = None
        attachment_filename = None
        
        # ВАЖНО: В БД PDF сохраняется в поле DOCX_CONTENT, а имя в DOCX_FILENAME
        if docx_content and docx_filename:
            try:
                # Читаем BLOB из Oracle
                if hasattr(docx_content, 'read'):
                    attachment_bytes = docx_content.read()
                else:
                    attachment_bytes = bytes(docx_content) if not isinstance(docx_content, bytes) else docx_content
                
                # Определяем тип файла по расширению имени файла
                filename_lower = docx_filename.lower()
                if filename_lower.endswith('.pdf'):
                    # Проверяем сигнатуру PDF
                    if len(attachment_bytes) >= 4 and attachment_bytes[:4] == b'%PDF':
                        attachment_content = attachment_bytes
                        attachment_filename = docx_filename
                        logging.info(f"Загружен PDF файл из БД: {attachment_filename}, размер: {len(attachment_bytes)} байт")
                    else:
                        logging.warning(f"Файл {docx_filename} имеет расширение .pdf, но не является PDF (сигнатура: {attachment_bytes[:4]})")
                        attachment_content = None
                        attachment_filename = None
                elif filename_lower.endswith('.docx'):
                    # DOCX используется только для извлечения текста, не как вложение
                    attachment_content = None
                    attachment_filename = None
                else:
                    # Если расширение неизвестно, проверяем по сигнатуре файла
                    if len(attachment_bytes) >= 4 and attachment_bytes[:4] == b'%PDF':
                        # Это PDF файл по сигнатуре
                        attachment_content = attachment_bytes
                        attachment_filename = docx_filename if docx_filename.endswith('.pdf') else docx_filename + '.pdf'
                        logging.info(f"Определен PDF файл по сигнатуре: {attachment_filename}")
                    else:
                        attachment_content = None
                        attachment_filename = None
            except Exception as e:
                logging.error(f"Ошибка при чтении вложения из БД: {e}")
                attachment_content = None
                attachment_filename = None
        
        # Формируем тело письма с приветствием (простой текст)
        email_body_text = greeting or 'Здравствуйте!'
        
        # Проверяем и убираем дублирование текста (используем единую функцию)
        original_length = len(email_body_text) if email_body_text else 0
        email_body_text = _remove_duplicate_text(email_body_text)
        
        if original_length > 0 and len(email_body_text) < original_length * 0.6:
            logging.info(f"Текст greeting был сокращен с {original_length} до {len(email_body_text)} символов (удалено дублирование)")
        
        # HTML версия: если текст уже содержит HTML-теги — используем как есть, иначе экранируем
        import html as html_module
        if _looks_like_html(email_body_text):
            html_body_content = email_body_text.strip()
        else:
            html_body_content = html_module.escape(email_body_text).replace('\n', '<br>').replace('\r', '')
        
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
    <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
        {html_body_content}
    </div>
</body>
</html>"""
        
        # Отправляем письма с задержками для избежания блокировки как спам
        sent_count = 0
        failed_count = 0
        error_messages = []
        
        # Вложения не используются в текущей реализации - отправляем напрямую через SMTP
        # Убрано создание временных файлов для избежания проблем с Untitled.bin
        
        # Отправляем письма с задержками для избежания блокировки как спам
        # Используем старую логику с добавленными задержками (работает без пароля)
        logging.info(f"Начало безопасной рассылки на {len(email_list)} адресов")
        logging.info(f"Задержка между письмами: {delay_between_emails} сек")
        logging.info(f"Задержка после партии ({batch_size} писем): {delay_after_batch} сек")
        
        try:
            # Подключение к SMTP серверу (как раньше, без обязательного пароля)
            server = smtplib.SMTP(smtp_host, smtp_port)
            if smtp_port == 587 or smtp_port == 465:
                server.starttls()  # Используем TLS если поддерживается
            
            # Аутентификация только если пароль указан
            if smtp_password:
                try:
                    server.login(from_email, smtp_password)
                except Exception as e:
                    logging.warning(f"Не удалось выполнить аутентификацию: {e}. Продолжаем без аутентификации.")
            
            for idx, email in enumerate(email_list):
                try:
                    # Создаем основное сообщение
                    # Если есть вложения, используем 'mixed', иначе просто HTML сообщение
                    if attachment_content and attachment_filename:
                        msg = MIMEMultipart('mixed')
                    else:
                        # Для простого HTML письма без вложений используем MIMEText напрямую
                        msg = MIMEText(full_html, 'html', 'utf-8')
                        msg['From'] = from_email
                        msg['To'] = email
                        msg['Subject'] = subject
                        # Отправляем сразу, без дополнительных attach
                        server.sendmail(from_email, [email], msg.as_string())
                        sent_count += 1
                        
                        # Задержка между письмами
                        if idx < len(email_list) - 1:
                            time.sleep(delay_between_emails)
                        
                        # Дополнительная задержка после каждой партии
                        if (idx + 1) % batch_size == 0 and idx < len(email_list) - 1:
                            logging.info(f"Отправлено {idx + 1}/{len(email_list)} писем. Пауза {delay_after_batch} сек...")
                            time.sleep(delay_after_batch)
                        continue
                    
                    # Если есть вложения, создаем multipart сообщение
                    msg['From'] = from_email
                    msg['To'] = email
                    msg['Subject'] = subject
                    
                    # Добавляем HTML версию
                    html_part = MIMEText(full_html, 'html', 'utf-8')
                    msg.attach(html_part)
                    
                    # Добавляем вложение (PDF)
                    if attachment_content and len(attachment_content) > 0:
                        # Проверяем, что это действительно PDF (по сигнатуре)
                        if attachment_content[:4] != b'%PDF':
                            logging.warning(f"Вложение не является PDF файлом (сигнатура: {attachment_content[:4]})")
                        
                        attachment_part = MIMEBase('application', 'pdf')
                        attachment_part.set_payload(attachment_content)
                        encoders.encode_base64(attachment_part)
                        
                        # Правильное формирование имени файла для вложения
                        from email.header import Header
                        from email.utils import encode_rfc2231
                        
                        # Обрабатываем имя файла для вложения
                        safe_filename = None
                        
                        if attachment_filename:
                            # Убираем путь, оставляем только имя файла
                            safe_filename = attachment_filename.split('/')[-1].split('\\')[-1]
                            
                            # Проверяем на некорректные имена
                            if (not safe_filename or 
                                safe_filename == "Untitled.bin" or 
                                safe_filename.lower().endswith('.bin') or
                                safe_filename == "attachment.pdf"):
                                safe_filename = None
                        
                        # Если имя файла некорректное или отсутствует, используем дефолтное
                        if not safe_filename:
                            safe_filename = "MVSAT_СТЭККОМ_26.pdf"  # Дефолтное имя для MVSAT рассылки
                            logging.warning(f"Имя файла не указано или некорректно ({attachment_filename}), используется дефолтное: {safe_filename}")
                        
                        # Убеждаемся, что имя файла заканчивается на .pdf
                        if not safe_filename.lower().endswith('.pdf'):
                            safe_filename = safe_filename.rsplit('.', 1)[0] + '.pdf'
                        
                        # Кодируем имя файла для поддержки кириллицы (RFC 2231)
                        try:
                            safe_filename.encode('ascii')
                            # Нет кириллицы, используем простое имя
                            filename_header = safe_filename
                        except UnicodeEncodeError:
                            # Есть кириллица, используем RFC 2231 кодирование
                            filename_header = encode_rfc2231(safe_filename, 'utf-8')
                        
                        # Устанавливаем заголовки правильно
                        attachment_part.add_header(
                            'Content-Disposition',
                            'attachment',
                            filename=filename_header
                        )
                        attachment_part.add_header('Content-Type', 'application/pdf')
                        attachment_part.add_header('Content-Transfer-Encoding', 'base64')
                        
                        msg.attach(attachment_part)
                        logging.info(f"Добавлено вложение PDF: {filename_header}, размер: {len(attachment_content)} байт, оригинальное имя в БД: {attachment_filename}")
                    else:
                        logging.warning("Вложение PDF пустое или отсутствует")
                    
                    # Отправляем
                    server.sendmail(from_email, [email], msg.as_string())
                    sent_count += 1
                    
                    # Задержка между письмами для избежания блокировки как спам
                    if idx < len(email_list) - 1:
                        time.sleep(delay_between_emails)
                    
                    # Дополнительная задержка после каждой партии
                    if (idx + 1) % batch_size == 0 and idx < len(email_list) - 1:
                        logging.info(f"Отправлено {idx + 1}/{len(email_list)} писем. Пауза {delay_after_batch} сек...")
                        time.sleep(delay_after_batch)
                except Exception as e:
                    failed_count += 1
                    error_messages.append(f"{email}: {str(e)}")
            
            server.quit()
            
            # Временные файлы больше не создаются - вложения отправляются напрямую
        except Exception as e:
            return 0, len(email_list), f"Ошибка SMTP подключения: {e}"
        
        # Обновляем статус кампании в БД
        if use_test_mode:
            status = 'TEST_SENT' if failed_count == 0 else ('PARTIAL' if sent_count > 0 else 'FAILED')
        else:
            status = 'SENT' if failed_count == 0 else ('PARTIAL' if sent_count > 0 else 'FAILED')
        error_msg = '; '.join(error_messages[:10]) if error_messages else None  # Первые 10 ошибок
        
        # Получаем текущего пользователя из параметра или из session_state
        if not sent_by:
            sent_by = st.session_state.get('username', 'unknown') if 'username' in st.session_state else 'system'
        
        cursor.execute("""
            UPDATE EMAIL_CAMPAIGNS
            SET STATUS = :1, EMAILS_SENT = :2, EMAILS_FAILED = :3,
                SENT_BY = :4, SENT_AT = SYSDATE, ERROR_MESSAGE = :5
            WHERE CAMPAIGN_ID = :6
        """, (status, sent_count, failed_count, sent_by, error_msg, campaign_id))
        
        conn.commit()
        cursor.close()
        
        result_msg = f"Отправлено: {sent_count}, Ошибок: {failed_count}"
        if error_messages:
            result_msg += f"\nПервые ошибки: {error_messages[:3]}"
        
        return sent_count, failed_count, result_msg
        
    except Exception as e:
        conn.rollback()
        return 0, 0, f"Ошибка при отправке кампании: {e}"


def _get_campaigns_list(conn, limit: int = 50) -> List[dict]:
    """Получить список кампаний из БД"""
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT CAMPAIGN_ID, CAMPAIGN_NAME, SUBJECT, STATUS,
                   EMAILS_TOTAL, EMAILS_SENT, EMAILS_FAILED,
                   CREATED_BY, CREATED_AT, SENT_AT, TEST_MODE, SENT_BY
            FROM EMAIL_CAMPAIGNS
            ORDER BY CREATED_AT DESC
            FETCH FIRST :1 ROWS ONLY
        """, (limit,))
        
        campaigns = []
        for row in cursor.fetchall():
            campaigns.append({
                'campaign_id': row[0],
                'campaign_name': row[1],
                'subject': row[2],
                'status': row[3],
                'emails_total': row[4] or 0,
                'emails_sent': row[5] or 0,
                'emails_failed': row[6] or 0,
                'created_by': row[7],
                'created_at': row[8],
                'sent_at': row[9],
                'test_mode': row[10] if len(row) > 10 else 0,
                'sent_by': row[11] if len(row) > 11 else None
            })
        
        cursor.close()
        return campaigns
    except Exception as e:
        st.error(f"❌ Ошибка при получении списка кампаний: {e}")
        return []


def _get_campaign_details(conn, campaign_id: int) -> Optional[dict]:
    """Получить детали кампании для повторного использования"""
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT CAMPAIGN_NAME, SUBJECT, GREETING, EMAIL_LIST,
                   DOCX_FILENAME, CREATED_BY, CREATED_AT, STATUS, TEST_MODE, TEST_EMAILS
            FROM EMAIL_CAMPAIGNS
            WHERE CAMPAIGN_ID = :1
        """, (campaign_id,))
        
        row = cursor.fetchone()
        if not row:
            return None
        
        def _lob_to_str(val, default=''):
            if val is None:
                return default
            if hasattr(val, 'read'):
                try:
                    return val.read()
                except Exception:
                    return str(val) if val else default
            return str(val) if val else default

        # Конвертируем LOB и обычные значения в строки (CLOB/GREETING и др.)
        campaign_name_val = _lob_to_str(row[0], '')
        if not isinstance(campaign_name_val, str):
            campaign_name_val = str(campaign_name_val) if campaign_name_val else ''
        subject_val = _lob_to_str(row[1], '')
        if not isinstance(subject_val, str):
            subject_val = str(subject_val) if subject_val else ''

        email_list_value = row[3]
        if email_list_value and hasattr(email_list_value, 'read'):
            try:
                email_list_value = email_list_value.read()
            except Exception:
                email_list_value = str(email_list_value) if email_list_value else ''
        else:
            email_list_value = email_list_value or ''

        greeting_val = row[2]
        if greeting_val is not None and hasattr(greeting_val, 'read'):
            try:
                greeting_val = greeting_val.read()
            except Exception:
                greeting_val = str(greeting_val) if greeting_val else ''
        else:
            greeting_val = greeting_val or ''
        
        test_emails_value = row[9] if len(row) > 9 else None
        if test_emails_value and hasattr(test_emails_value, 'read'):
            try:
                test_emails_value = test_emails_value.read()
            except Exception:
                test_emails_value = str(test_emails_value) if test_emails_value else None
        
        return {
            'campaign_name': campaign_name_val,
            'subject': subject_val,
            'greeting': greeting_val,
            'email_list': email_list_value,
            'docx_filename': row[4],
            'created_by': row[5],
            'created_at': row[6],
            'status': row[7],
            'test_mode': row[8] if len(row) > 8 else 0,
            'test_emails': test_emails_value
        }
    except Exception as e:
        st.error(f"❌ Ошибка при получении деталей кампании: {e}")
        return None


def _update_campaign_metadata(conn, campaign_id: int, campaign_name: str, subject: str, greeting: str) -> bool:
    """Обновить название, тему и тело письма (GREETING) существующей кампании."""
    try:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE EMAIL_CAMPAIGNS
            SET CAMPAIGN_NAME = :1, SUBJECT = :2, GREETING = :3
            WHERE CAMPAIGN_ID = :4
        """, (campaign_name or '', subject or '', greeting or '', campaign_id))
        conn.commit()
        updated = cursor.rowcount
        cursor.close()
        return updated > 0
    except Exception as e:
        st.error(f"❌ Ошибка при обновлении кампании: {e}")
        if conn:
            conn.rollback()
        return False


def show_tab():
    """
    Закладка для создания и отправки рекламных email кампаний
    """
    st.header("📧 Рекламные email кампании")
    st.markdown(
        """
        Эта вкладка позволяет создавать и отправлять рекламные email кампании клиентам.
        
        **Как использовать:**
        1. Введите тему письма и текст приветствия
        2. Загрузите текстовый файл со списком email (разделитель - запятая) или введите вручную
        3. (Опционально) Загрузите PDF файл для вложения
        4. Сохраните кампанию и отправьте письма
        5. Используйте сохраненные кампании для повторной отправки
        """
    )
    
    # Проверка подключения к БД
    conn = _get_db_connection()
    if not conn:
        st.error("❌ Не удалось подключиться к базе данных Oracle")
        st.info("Проверьте настройки подключения в config.env или переменные окружения")
        return
    
    # Проверка существования таблицы
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT COUNT(*) FROM user_tables WHERE table_name = 'EMAIL_CAMPAIGNS'
        """)
        table_exists = cursor.fetchone()[0] > 0
        cursor.close()
        
        if not table_exists:
            st.warning("⚠️ Таблица EMAIL_CAMPAIGNS не найдена в базе данных.")
            st.info("Выполните SQL скрипт: `oracle/tables/05_email_campaigns.sql`")
            conn.close()
            return
    except Exception as e:
        st.error(f"❌ Ошибка при проверке таблицы: {e}")
        conn.close()
        return
    
    # Получаем текущего пользователя
    username = st.session_state.get('username', 'unknown')
    
    # Вкладки для разных функций
    tab_new, tab_list, tab_reuse = st.tabs(["➕ Новая кампания", "📋 Список кампаний", "♻️ Повторное использование"])
    
    with tab_new:
        st.subheader("Создание новой кампании")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            campaign_name = st.text_input(
                "Название кампании",
                value=f"Кампания {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                help="Уникальное название для идентификации кампании"
            )
            
            subject = st.text_input(
                "Тема письма",
                placeholder="Например: Специальное предложение от STECCOM",
                help="Тема email письма",
                key="subject_input"
            )
            
            greeting = st.text_area(
                "Текст письма (поддержка HTML)",
                placeholder="Введите текст или HTML: <b>жирный</b>, <span style='color:red'>цвет</span>, <br> — перенос",
                help="Можно вводить обычный текст или HTML: <b>, <i>, <span style='color:...'>, <font>, переносы сохраняются.",
                key="greeting_input",
                height=220
            )
            if greeting.strip():
                with st.expander("👁 Предпросмотр письма"):
                    if _looks_like_html(greeting):
                        st.markdown(greeting, unsafe_allow_html=True)
                    else:
                        st.markdown(greeting.replace("\n", "  \n"))
            
            # Проверяем дублирование текста в интерфейсе
            if greeting and len(greeting) > 20:
                greeting_stripped = greeting.strip()
                text_length = len(greeting_stripped)
                half_length = text_length // 2
                first_half = greeting_stripped[:half_length].strip()
                second_half = greeting_stripped[half_length:].strip()
                if first_half == second_half and len(first_half) > 10:
                    st.warning("⚠️ **Обнаружено дублирование текста!** Текст повторяется дважды. Дублирование будет автоматически удалено при отправке.")
                    st.info(f"Первая часть текста (будет использована):\n\n{first_half[:200]}...")
        
        with col2:
            from_email = st.text_input(
                "Обратный адрес",
                value="sales@steccom.ru",
                help="Email адрес отправителя"
            )
            
            smtp_host = st.text_input(
                "SMTP сервер",
                value="mail.steccom.ru",
                help="Адрес SMTP сервера"
            )
            
            smtp_port = st.number_input(
                "SMTP порт",
                value=25,
                min_value=1,
                max_value=65535,
                help="Порт SMTP сервера"
            )
            
            smtp_password = st.text_input(
                "Пароль SMTP (опционально)",
                type="password",
                help="Пароль для аутентификации на SMTP сервере. Оставьте пустым, если сервер не требует аутентификации.",
                key="smtp_password_input"
            )
        
        st.markdown("---")
        
        # Настройки безопасной рассылки
        with st.expander("⚙️ Настройки безопасной рассылки (для избежания блокировки как спам)", expanded=False):
            st.markdown("""
            **Рекомендации:**
            - Задержка между письмами: 2-5 секунд для небольших рассылок, 5-10 секунд для больших
            - Размер партии: 10-20 писем перед длительной паузой
            - Задержка после партии: 60-120 секунд для избежания блокировки
            """)
            
            col_delay1, col_delay2, col_delay3 = st.columns(3)
            
            with col_delay1:
                delay_between_emails = st.number_input(
                    "Задержка между письмами (сек)",
                    value=2.0,
                    min_value=0.0,
                    max_value=60.0,
                    step=0.5,
                    help="Пауза между отправкой каждого письма"
                )
            
            with col_delay2:
                batch_size = st.number_input(
                    "Размер партии",
                    value=10,
                    min_value=1,
                    max_value=100,
                    help="Количество писем перед длительной паузой"
                )
            
            with col_delay3:
                delay_after_batch = st.number_input(
                    "Задержка после партии (сек)",
                    value=60.0,
                    min_value=0.0,
                    max_value=600.0,
                    step=5.0,
                    help="Длительная пауза после каждой партии писем"
                )
        
        st.markdown("---")
        
        # Загрузка списка email
        st.subheader("📧 Список получателей")
        
        uploaded_email_file = st.file_uploader(
            "Загрузите текстовый файл со списком email",
            type=['txt', 'csv'],
            help="Файл должен содержать email адреса, разделенные запятыми, точкой с запятой или переносами строк"
        )
        
        email_file = uploaded_email_file if uploaded_email_file else None
        
        email_text_input = st.text_area(
            "Или введите email адреса вручную",
            placeholder="email1@example.com, email2@example.com, email3@example.com",
            help="Введите email адреса через запятую, точку с запятой или с новой строки"
        )
        
        email_list = []
        if email_file:
            try:
                email_payload = None
                # Для Streamlit UploadedFile используем getvalue(), чтобы не "съесть" файл при предпросмотре
                if hasattr(email_file, "getvalue"):
                    email_payload = email_file.getvalue()
                elif hasattr(email_file, "read"):
                    email_payload = email_file.read()

                if isinstance(email_payload, bytes):
                    email_text = email_payload.decode("utf-8", errors="replace")
                elif email_payload is not None:
                    email_text = str(email_payload)
                else:
                    email_text = str(email_file)

                email_list = _parse_email_list(email_text)
                st.info(f"✅ Загружено {len(email_list)} валидных email адресов из файла")
            except Exception as e:
                st.error(f"❌ Ошибка при чтении файла email: {e}")
        elif email_text_input:
            email_list = _parse_email_list(email_text_input)
            if email_list:
                st.info(f"✅ Найдено {len(email_list)} валидных email адресов")
        
        if email_list:
            with st.expander(f"Просмотр списка получателей ({len(email_list)} адресов)"):
                st.write(', '.join(email_list[:50]))
                if len(email_list) > 50:
                    st.write(f"... и еще {len(email_list) - 50} адресов")
        
        st.markdown("---")
        
        # Загрузка вложения PDF
        st.subheader("📎 Вложение (PDF)")
        st.caption("Опционально: загрузите PDF файл, который будет прикреплен к письму")
        
        uploaded_pdf_file = st.file_uploader(
            "Загрузите PDF файл для вложения",
            type=['pdf'],
            help="PDF файл будет прикреплен к каждому письму",
            key="pdf_uploader"
        )
        
        pdf_file = uploaded_pdf_file
        
        if pdf_file:
            try:
                if hasattr(pdf_file, "getvalue"):
                    file_size = len(pdf_file.getvalue() or b"")
                elif hasattr(pdf_file, "read"):
                    content = pdf_file.read()
                    file_size = len(content) if isinstance(content, bytes) else 0
                else:
                    file_size = 0
                file_name = getattr(pdf_file, "name", "unknown")
                st.success(f"✅ {file_name} ({file_size} байт)")
            except Exception as e:
                st.warning(f"⚠️ Не удалось определить размер файла: {e}")
        
        st.markdown("---")
        
        # Контрольные email для тестовой рассылки
        st.subheader("🧪 Контрольные email для тестовой рассылки")
        st.markdown(
            """
            **Важно:** Укажите контрольные email адреса для проверки рассылки перед отправкой основному списку.
            Эти адреса будут использоваться при тестовой рассылке.
            """
        )
        
        test_emails_input = st.text_area(
            "Контрольные email адреса",
            placeholder="test1@example.com, test2@example.com",
            help="Введите контрольные email адреса через запятую. Они будут использоваться для тестовой рассылки.",
            key="test_emails_input"
        )
        
        test_emails_list = []
        if test_emails_input:
            test_emails_list = _parse_email_list(test_emails_input)
            if test_emails_list:
                st.info(f"✅ Найдено {len(test_emails_list)} валидных контрольных email адресов")
                st.write(', '.join(test_emails_list))
            else:
                st.warning("⚠️ Не найдено валидных email адресов в поле контрольных email")
        
        st.markdown("---")
        
        # Режим рассылки
        st.subheader("📤 Режим рассылки")
        test_mode = st.checkbox(
            "🧪 Тестовая рассылка",
            value=False,
            help="Если включено, рассылка будет отправлена только на контрольные email адреса, указанные выше. "
                 "Основной список получателей будет проигнорирован.",
            key="test_mode_checkbox"
        )
        
        if test_mode:
            if not test_emails_list:
                st.error("❌ Для тестовой рассылки необходимо указать хотя бы один контрольный email адрес!")
            else:
                st.warning(f"⚠️ **РЕЖИМ ТЕСТОВОЙ РАССЫЛКИ АКТИВЕН**")
                st.info(f"📧 Письма будут отправлены только на {len(test_emails_list)} контрольных адресов:")
                st.write(', '.join(test_emails_list))
                
                # Показываем основной список для информации
                if email_list:
                    st.markdown("---")
                    st.subheader("📋 Основной список получателей (для боевой рассылки)")
                    st.info(f"📬 При обычной рассылке письма будут отправлены на {len(email_list)} адресов из основного списка:")
                    with st.expander(f"Просмотр основного списка ({len(email_list)} адресов)", expanded=False):
                        st.write(', '.join(email_list[:100]))
                        if len(email_list) > 100:
                            st.write(f"... и еще {len(email_list) - 100} адресов")
                    st.caption("💡 Этот список будет использован при обычной (боевой) рассылке после проверки тестовой")
                else:
                    st.warning("⚠️ Основной список получателей пуст. При обычной рассылке не будет получателей!")
        else:
            if email_list:
                st.info(f"📧 Обычная рассылка: письма будут отправлены на {len(email_list)} адресов из основного списка")
                with st.expander(f"Просмотр списка получателей ({len(email_list)} адресов)", expanded=False):
                    st.write(', '.join(email_list[:100]))
                    if len(email_list) > 100:
                        st.write(f"... и еще {len(email_list) - 100} адресов")
            else:
                st.warning("⚠️ Основной список получателей пуст")
        
        st.markdown("---")
        
        # Кнопки действий
        col_save, col_send = st.columns(2)
        
        with col_save:
            if st.button("💾 Сохранить кампанию", type="primary", use_container_width=True):
                if not campaign_name:
                    st.error("Введите название кампании")
                elif not subject:
                    st.error("Введите тему письма")
                elif not test_mode and not email_list:
                    # Для обычной рассылки список обязателен, для тестовой - может быть пустым
                    st.error("Загрузите список email получателей")
                elif test_mode and not test_emails_list:
                    # Для тестовой рассылки нужны тестовые email
                    st.error("Для тестовой рассылки укажите контрольные email адреса")
                else:
                    with st.spinner("Сохранение кампании..."):
                        # Сохраняем PDF как вложение (в BLOB)
                        attachment_content = None
                        attachment_filename = None
                        if pdf_file:
                            try:
                                if hasattr(pdf_file, "getvalue"):
                                    attachment_content = pdf_file.getvalue()
                                elif hasattr(pdf_file, "read"):
                                    attachment_content = pdf_file.read()
                                
                                # Получаем имя файла и проверяем его корректность
                                raw_filename = getattr(pdf_file, "name", "attachment.pdf")
                                
                                # Если имя файла некорректное (Untitled.bin, пустое и т.д.), используем дефолтное
                                if not raw_filename or raw_filename == "Untitled.bin" or raw_filename.lower().endswith('.bin'):
                                    # Пытаемся определить имя из пути по умолчанию или используем стандартное
                                    if default_pdf_file_path and default_pdf_file_path.exists():
                                        attachment_filename = default_pdf_file_path.name
                                    else:
                                        attachment_filename = "MVSAT_СТЭККОМ_26.pdf"
                                    logging.warning(f"Имя файла некорректное ({raw_filename}), используется: {attachment_filename}")
                                else:
                                    # Убираем путь, оставляем только имя файла
                                    attachment_filename = raw_filename.split('/')[-1].split('\\')[-1]
                                    
                                    # Если имя не заканчивается на .pdf, добавляем расширение
                                    if not attachment_filename.lower().endswith('.pdf'):
                                        attachment_filename = attachment_filename + '.pdf'
                                    
                            except Exception as e:
                                st.warning(f"⚠️ Ошибка при чтении PDF файла: {e}")
                        
                        campaign_id = _save_campaign_to_db(
                            conn,
                            campaign_name,
                            subject,
                            greeting,
                            email_list,
                            attachment_content,  # PDF сохраняется в BLOB
                            attachment_filename,  # Имя PDF файла
                            username,
                            test_mode,
                            test_emails_list if test_emails_list else None
                        )
                        if campaign_id:
                            mode_text = "тестовой" if test_mode else "обычной"
                            st.success(f"✅ Кампания сохранена! ID: {campaign_id} ({mode_text} рассылки)")
                            st.info("Теперь вы можете отправить кампанию или использовать её позже")
                        else:
                            st.error("Не удалось сохранить кампанию")
        
        with col_send:
            if st.button("📤 Сохранить и отправить", type="primary", use_container_width=True):
                if not campaign_name:
                    st.error("Введите название кампании")
                elif not subject:
                    st.error("Введите тему письма")
                elif test_mode and not test_emails_list:
                    st.error("Для тестовой рассылки необходимо указать контрольные email адреса")
                elif not test_mode and not email_list:
                    st.error("Загрузите список email получателей")
                else:
                    with st.spinner("Сохранение и отправка кампании..."):
                        # Сначала сохраняем PDF как вложение (в BLOB)
                        attachment_content = None
                        attachment_filename = None
                        if pdf_file:
                            try:
                                if hasattr(pdf_file, "getvalue"):
                                    attachment_content = pdf_file.getvalue()
                                elif hasattr(pdf_file, "read"):
                                    attachment_content = pdf_file.read()
                                
                                # Получаем имя файла и проверяем его корректность
                                raw_filename = getattr(pdf_file, "name", "attachment.pdf")
                                
                                # Если имя файла некорректное (Untitled.bin, пустое и т.д.), используем дефолтное
                                if not raw_filename or raw_filename == "Untitled.bin" or raw_filename.lower().endswith('.bin'):
                                    # Пытаемся определить имя из пути по умолчанию или используем стандартное
                                    if default_pdf_file_path and default_pdf_file_path.exists():
                                        attachment_filename = default_pdf_file_path.name
                                    else:
                                        attachment_filename = "MVSAT_СТЭККОМ_26.pdf"
                                    logging.warning(f"Имя файла некорректное ({raw_filename}), используется: {attachment_filename}")
                                else:
                                    # Убираем путь, оставляем только имя файла
                                    attachment_filename = raw_filename.split('/')[-1].split('\\')[-1]
                                    
                                    # Если имя не заканчивается на .pdf, добавляем расширение
                                    if not attachment_filename.lower().endswith('.pdf'):
                                        attachment_filename = attachment_filename + '.pdf'
                                    
                            except Exception as e:
                                st.warning(f"⚠️ Ошибка при чтении PDF файла: {e}")
                        
                        campaign_id = _save_campaign_to_db(
                            conn,
                            campaign_name,
                            subject,
                            greeting,
                            email_list,
                            attachment_content,  # PDF сохраняется в BLOB
                            attachment_filename,  # Имя PDF файла
                            username,
                            test_mode,
                            test_emails_list if test_emails_list else None
                        )
                        
                        if campaign_id:
                            st.success(f"✅ Кампания сохранена (ID: {campaign_id})")
                            
                            # Затем отправляем
                            recipients_count = len(test_emails_list) if test_mode else len(email_list)
                            with st.spinner(f"Отправка писем {recipients_count} получателям (с задержками для безопасности)..."):
                                sent, failed, msg = _send_email_campaign(
                                    conn,
                                    campaign_id,
                                    smtp_host,
                                    int(smtp_port),
                                    from_email,
                                    smtp_password if smtp_password else None,
                                    test_mode,
                                    test_emails_list if test_emails_list else None,
                                    username,
                                    delay_between_emails,
                                    delay_after_batch,
                                    batch_size
                                )
                                
                                if sent > 0:
                                    st.success(f"✅ {msg}")
                                elif failed > 0:
                                    st.warning(f"⚠️ {msg}")
                                else:
                                    st.error(f"❌ {msg}")
                        else:
                            st.error("Не удалось сохранить кампанию")
    
    with tab_list:
        st.subheader("Список всех кампаний")
        
        campaigns = _get_campaigns_list(conn, limit=100)
        
        if not campaigns:
            st.info("Пока нет сохраненных кампаний")
        else:
            import pandas as pd
            
            df_data = []
            for camp in campaigns:
                status_display = camp['status']
                if camp['test_mode']:
                    status_display = f"🧪 {status_display}"
                df_data.append({
                    'ID': camp['campaign_id'],
                    'Название': camp['campaign_name'],
                    'Тема': camp['subject'],
                    'Статус': status_display,
                    'Режим': 'Тест' if camp['test_mode'] else 'Обычный',
                    'Всего': camp['emails_total'],
                    'Отправлено': camp['emails_sent'],
                    'Ошибок': camp['emails_failed'],
                    'Создал': camp['created_by'],
                    'Создано': camp['created_at'].strftime('%Y-%m-%d %H:%M') if camp['created_at'] else '-',
                    'Отправил': camp['sent_by'] or '-',
                    'Отправлено': camp['sent_at'].strftime('%Y-%m-%d %H:%M') if camp['sent_at'] else '-'
                })
            
            df = pd.DataFrame(df_data)
            st.dataframe(df, use_container_width=True, hide_index=True)
    
    with tab_reuse:
        st.subheader("Повторное использование сохраненной кампании")
        
        campaigns = _get_campaigns_list(conn, limit=50)
        
        if not campaigns:
            st.info("Нет сохраненных кампаний для повторного использования")
        else:
            campaign_options = {f"{c['campaign_id']}: {c['campaign_name']}": c['campaign_id'] 
                               for c in campaigns}
            
            selected_campaign_key = st.selectbox(
                "Выберите кампанию",
                options=list(campaign_options.keys()),
                help="Выберите сохраненную кампанию для повторной отправки"
            )
            
            if selected_campaign_key:
                campaign_id = campaign_options[selected_campaign_key]
                details = _get_campaign_details(conn, campaign_id)
                
                if details:
                    st.markdown("---")
                    st.subheader("Детали кампании")
                    
                    col_d1, col_d2 = st.columns(2)
                    with col_d1:
                        st.write(f"**Название:** {details['campaign_name']}")
                        st.write(f"**Тема:** {details['subject']}")
                        st.write(f"**Приветствие:** {details['greeting']}")
                    with col_d2:
                        st.write(f"**Статус:** {details['status']}")
                        st.write(f"**Создал:** {details['created_by']}")
                        st.write(f"**Создано:** {details['created_at'].strftime('%Y-%m-%d %H:%M') if details['created_at'] else '-'}")
                    
                    st.markdown("---")
                    st.subheader("Список получателей")
                    email_list_reuse = _parse_email_list(details['email_list'])
                    test_emails_reuse = _parse_email_list(details['test_emails']) if details.get('test_emails') else []
                    
                    if details.get('test_mode'):
                        st.warning("🧪 **Это была тестовая кампания**")
                        if test_emails_reuse:
                            st.write(f"Контрольные email: {len(test_emails_reuse)}")
                            with st.expander("Просмотр контрольных email"):
                                st.write(', '.join(test_emails_reuse))
                    else:
                        st.write(f"Основной список: {len(email_list_reuse)} email")
                        with st.expander("Просмотр списка"):
                            st.write(', '.join(email_list_reuse[:50]))
                            if len(email_list_reuse) > 50:
                                st.write(f"... и еще {len(email_list_reuse) - 50} адресов")
                    
                    st.markdown("---")
                    st.subheader("Повторная отправка")
                    
                    test_mode_reuse = st.checkbox(
                        "🧪 Тестовая рассылка",
                        value=details.get('test_mode', False),
                        help="Если включено, рассылка будет отправлена только на контрольные email адреса",
                        key="reuse_test_mode"
                    )
                    
                    test_emails_reuse_input = st.text_area(
                        "Контрольные email (для тестовой рассылки)",
                        value=details.get('test_emails', '') or '',
                        placeholder="test1@example.com, test2@example.com",
                        key="reuse_test_emails"
                    )
                    
                    test_emails_reuse_parsed = []
                    if test_emails_reuse_input:
                        test_emails_reuse_parsed = _parse_email_list(test_emails_reuse_input)
                    
                    col_r1, col_r2 = st.columns(2)
                    with col_r1:
                        smtp_host_reuse = st.text_input("SMTP сервер", value="mail.steccom.ru", key="reuse_smtp_host")
                        smtp_port_reuse = st.number_input("SMTP порт", value=25, min_value=1, max_value=65535, key="reuse_smtp_port")
                        smtp_password_reuse = st.text_input(
                            "Пароль SMTP (опционально)",
                            type="password",
                            help="Пароль для аутентификации на SMTP сервере. Оставьте пустым, если сервер не требует аутентификации.",
                            key="reuse_smtp_password"
                        )
                    with col_r2:
                        from_email_reuse = st.text_input("Обратный адрес", value="sales@steccom.ru", key="reuse_from_email")
                    
                    # Настройки безопасной рассылки для повторной отправки
                    with st.expander("⚙️ Настройки безопасной рассылки", expanded=False):
                        col_delay_r1, col_delay_r2, col_delay_r3 = st.columns(3)
                        with col_delay_r1:
                            delay_between_emails_reuse = st.number_input(
                                "Задержка между письмами (сек)",
                                value=2.0,
                                min_value=0.0,
                                max_value=60.0,
                                step=0.5,
                                key="reuse_delay_between"
                            )
                        with col_delay_r2:
                            batch_size_reuse = st.number_input(
                                "Размер партии",
                                value=10,
                                min_value=1,
                                max_value=100,
                                key="reuse_batch_size"
                            )
                        with col_delay_r3:
                            delay_after_batch_reuse = st.number_input(
                                "Задержка после партии (сек)",
                                value=60.0,
                                min_value=0.0,
                                max_value=600.0,
                                step=5.0,
                                key="reuse_delay_after_batch"
                            )
                    
                    if st.button("📤 Отправить кампанию повторно", type="primary", use_container_width=True):
                        if test_mode_reuse and not test_emails_reuse_parsed:
                            st.error("Для тестовой рассылки необходимо указать контрольные email адреса")
                        else:
                            recipients_count = len(test_emails_reuse_parsed) if test_mode_reuse else len(email_list_reuse)
                            mode_text = "тестовой" if test_mode_reuse else "обычной"
                            with st.spinner(f"Отправка писем ({mode_text} режим) {recipients_count} получателям (с задержками для безопасности)..."):
                                sent, failed, msg = _send_email_campaign(
                                    conn,
                                    campaign_id,
                                    smtp_host_reuse,
                                    int(smtp_port_reuse),
                                    from_email_reuse,
                                    smtp_password_reuse if smtp_password_reuse else None,
                                    test_mode_reuse,
                                    test_emails_reuse_parsed if test_emails_reuse_parsed else None,
                                    username,
                                    delay_between_emails_reuse,
                                    delay_after_batch_reuse,
                                    batch_size_reuse
                                )
                            
                            if sent > 0:
                                st.success(f"✅ {msg}")
                            elif failed > 0:
                                st.warning(f"⚠️ {msg}")
                            else:
                                st.error(f"❌ {msg}")
    
    conn.close()

