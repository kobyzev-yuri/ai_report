#!/usr/bin/env python3
"""
Парсеры вложений Confluence: интеллектуальная обработка через Gemini (без OCR).
- PDF: текст страниц + картинки в PDF описываются через Gemini Vision.
- Draw.io: из формата извлекается структура, описание и смысловые блоки — через Gemini.
- Изображения: только Gemini Vision (описание и смысловые блоки).
- DOCX: извлечение текста как прежде.
"""
import base64
import io
import logging
import re
import zlib
import xml.etree.ElementTree as ET
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# Расширения и MIME для выбора парсера
PDF_EXT = (".pdf",)
PDF_MIME = ("application/pdf",)
DOCX_EXT = (".docx", ".doc")
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
)
XLS_EXT = (".xls",)
XLS_MIME = ("application/vnd.ms-excel",)
# .xlsx — опционально, через openpyxl
XLSX_EXT = (".xlsx",)
XLSX_MIME = ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",)
DRAWIO_EXT = (".drawio", ".drawio.xml", ".xml")
# Draw.io в Confluence часто как application/xml или octet-stream
IMAGE_EXT = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp")
IMAGE_MIME = ("image/png", "image/jpeg", "image/gif", "image/bmp", "image/tiff", "image/webp")


def _normalize_ext(filename: str) -> str:
    if not filename:
        return ""
    p = filename.rfind(".")
    return filename[p:].lower() if p >= 0 else ""


def extract_text_from_pdf(
    data: bytes,
    filename: str = "",
    context_text: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    PDF: извлечь текст страниц и картинки. Картинки описываются через Gemini Vision;
    текст страниц остаётся как есть.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "", "PyMuPDF не установлен: pip install pymupdf"
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        parts: list = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text().strip()
            if page_text:
                parts.append("Стр. %s:\n%s" % (page_num + 1, page_text))
            # Картинки на странице — описание через Gemini Vision
            try:
                from kb_billing.rag.vision_annotation import describe_image, is_vision_available
                if is_vision_available():
                    for xref in page.get_images():
                        if not isinstance(xref, (list, tuple)):
                            continue
                        img_xref = xref[0]
                        try:
                            img_info = doc.extract_image(img_xref)
                            img_bytes = img_info.get("image")
                            ext = (img_info.get("ext") or "png").lower()
                            mime = "image/png" if ext == "png" else "image/jpeg"
                            if ext == "webp":
                                mime = "image/webp"
                            if img_bytes:
                                ctx = (context_text or "") + " PDF: %s, стр. %s" % (filename or "?", page_num + 1)
                                desc = describe_image(img_bytes, mime_type=mime, context_text=ctx.strip())
                                if desc:
                                    parts.append("Изображение на стр. %s:\n%s" % (page_num + 1, desc))
                        except Exception as e:
                            logger.debug("Извлечение изображения из PDF: %s", e)
            except Exception as e:
                logger.warning("Vision для PDF-изображений не удался: %s", e)
        doc.close()
        text = "\n\n".join(parts) if parts else "(документ без текста и без описанных изображений)"
        return text.strip(), None
    except Exception as e:
        logger.warning("PDF parse %s: %s", filename or "?", e)
        return "", str(e)


def extract_text_from_xls(data: bytes, filename: str = "") -> Tuple[str, Optional[str]]:
    """
    Извлечь текст из XLS (Excel 97–2003): названия листов и содержимое ячеек.
    Для контекста в KB — списки устройств, конфигурации, таблицы.
    """
    try:
        import xlrd
    except ImportError:
        return "", "xlrd не установлен: pip install xlrd"
    try:
        book = xlrd.open_workbook(file_contents=data)
    except Exception as e:
        logger.warning("XLS open %s: %s", filename or "?", e)
        return "", str(e)
    parts = []
    max_cells = 50000  # ограничение на число ячеек, чтобы не раздувать чанк
    total_cells = 0
    for sheet_idx in range(book.nsheets):
        sheet = book.sheet_by_index(sheet_idx)
        name = sheet.name or f"Лист{sheet_idx + 1}"
        parts.append(f"Лист: {name}")
        for row_idx in range(sheet.nrows):
            if total_cells >= max_cells:
                parts.append("[... таблица обрезана по лимиту ячеек ...]")
                break
            row_vals = []
            for col_idx in range(sheet.ncols):
                cell = sheet.cell(row_idx, col_idx)
                if cell.ctype == xlrd.XL_CELL_EMPTY:
                    row_vals.append("")
                elif cell.ctype == xlrd.XL_CELL_TEXT:
                    row_vals.append(cell.value.strip() if cell.value else "")
                elif cell.ctype == xlrd.XL_CELL_NUMBER:
                    if cell.value == int(cell.value):
                        row_vals.append(str(int(cell.value)))
                    else:
                        row_vals.append(str(cell.value))
                elif cell.ctype == xlrd.XL_CELL_DATE:
                    row_vals.append(str(cell.value))
                else:
                    row_vals.append(str(cell.value) if cell.value else "")
                total_cells += 1
            line = " | ".join(row_vals).strip()
            if line:
                parts.append(line)
        if total_cells >= max_cells:
            break
    text = "\n".join(parts).strip() or "(таблица пуста)"
    return text, None


def extract_text_from_xlsx(data: bytes, filename: str = "") -> Tuple[str, Optional[str]]:
    """
    Извлечь текст из XLSX (Excel 2007+). Аналогично XLS — листы и ячейки для контекста KB.
    """
    try:
        import openpyxl
    except ImportError:
        return "", "openpyxl не установлен: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:
        logger.warning("XLSX open %s: %s", filename or "?", e)
        return "", str(e)
    parts = []
    max_cells = 50000
    total_cells = 0
    for name in wb.sheetnames:
        parts.append(f"Лист: {name}")
        sheet = wb[name]
        for row in sheet.iter_rows(values_only=True):
            if total_cells >= max_cells:
                parts.append("[... таблица обрезана по лимиту ячеек ...]")
                break
            row_vals = [str(c) if c is not None else "" for c in row]
            line = " | ".join(row_vals).strip()
            if line:
                parts.append(line)
            total_cells += len(row_vals)
        if total_cells >= max_cells:
            break
    wb.close()
    text = "\n".join(parts).strip() or "(таблица пуста)"
    return text, None


def extract_text_from_docx(data: bytes, filename: str = "") -> Tuple[str, Optional[str]]:
    """
    Извлечь текст из DOCX. Возвращает (plain_text, error_message).
    """
    try:
        from docx import Document
    except ImportError:
        return "", "python-docx не установлен: pip install python-docx"
    try:
        doc = Document(io.BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        return text.strip() or "(документ без текста)", None
    except Exception as e:
        logger.warning("DOCX parse %s: %s", filename or "?", e)
        return "", str(e)


def _strip_html_fragment(html: str) -> str:
    """Убрать теги из фрагмента HTML (value в draw.io может быть <p>...</p>)."""
    if not html:
        return ""
    return re.sub(r"<[^>]+>", " ", html).replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").strip()


def _drawio_decompress(diagram_content: str) -> Optional[str]:
    """Распаковать сжатый контент из <diagram> (base64 + raw deflate)."""
    s = (diagram_content or "").strip()
    if not s or s.startswith("<"):
        return s or None
    try:
        raw = base64.b64decode(s)
        # raw deflate без zlib-заголовка
        return zlib.decompress(raw, -zlib.MAX_WBITS).decode("utf-8", errors="replace")
    except Exception:
        return None


def extract_text_from_drawio(
    data: bytes,
    filename: str = "",
    context_text: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    Draw.io: извлечь структуру (подписи, узлы) из формата, затем описание и смысловые блоки через Gemini.
    Если Gemini недоступен — возвращается сырой извлечённый текст.
    """
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception as e:
        return "", f"Ошибка декодирования: {e}"
    text = text.strip()
    if not text.startswith("<"):
        return "", "Файл не похож на XML (draw.io: сохраните без сжатия в настройках)"
    try:
        root = ET.fromstring(text)
    except ET.ParseError as e:
        return "", f"Ошибка разбора XML: {e}"

    # Сжатый draw.io: mxfile > diagram с base64+deflate внутри
    tag = root.tag.split("}")[-1] if "}" in root.tag else root.tag
    if tag == "mxfile":

        def find_diagram(elem: ET.Element):
            if elem.tag.split("}")[-1] == "diagram":
                return elem
            for c in elem:
                found = find_diagram(c)
                if found is not None:
                    return found
            return None

        diagram = find_diagram(root)
        if diagram is not None:
            raw = (diagram.text or "").strip()
            if raw and not raw.startswith("<"):
                inner = _drawio_decompress(raw)
                if inner:
                    try:
                        root = ET.fromstring(inner)
                    except ET.ParseError:
                        pass
            elif raw and raw.startswith("<"):
                try:
                    root = ET.fromstring(raw)
                except ET.ParseError:
                    pass

    texts: list[str] = []

    def collect_mxcell(elem: ET.Element) -> None:
        t = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if t == "mxCell":
            val = elem.get("value") or ""
            if val:
                texts.append(_strip_html_fragment(val))
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
        for child in elem:
            collect_mxcell(child)
            if child.tail and child.tail.strip():
                texts.append(child.tail.strip())

    collect_mxcell(root)
    seen = set()
    unique = []
    for t in texts:
        t = t.strip()
        if t and t not in seen:
            seen.add(t)
            unique.append(t)
    result = "\n".join(unique)
    raw_text = result or "(схема без текстовых подписей)"
    try:
        from kb_billing.rag.vision_annotation import describe_drawio_content, is_vision_available
        if is_vision_available():
            description = describe_drawio_content(raw_text, context_text=context_text)
            if description:
                return description, None
    except Exception as e:
        logger.warning("Gemini описание draw.io не удалось: %s", e)
    return raw_text, None


def _mime_for_image(filename: str, data: bytes) -> str:
    """Определить MIME по имени или по magic bytes."""
    ext = (filename or "").lower()
    if ext.endswith(".png"):
        return "image/png"
    if ext.endswith(".webp"):
        return "image/webp"
    if ext.endswith(".gif"):
        return "image/gif"
    if ext.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if ext.endswith((".tiff", ".tif")):
        return "image/tiff"
    if len(data) >= 8 and data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if len(data) >= 2 and data[:2] in (b"\xff\xd8", b"\xFF\xD8"):
        return "image/jpeg"
    return "image/png"


def extract_text_from_image(
    data: bytes,
    filename: str = "",
    context_text: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    Интеллектуальное описание изображения/схемы только через Gemini Vision (без OCR).
    Возвращает описание и смысловые блоки или placeholder при недоступности Vision.
    """
    try:
        from kb_billing.rag.vision_annotation import describe_image, is_vision_available
        if not is_vision_available():
            return "[Изображение/схема — задайте GEMINI_API_KEY или OPENAI_API_KEY в config.env для аннотации через Gemini Vision]", None
        mime = _mime_for_image(filename, data)
        description = describe_image(
            data,
            mime_type=mime,
            filename=filename,
            context_text=context_text,
        )
        if description:
            return description, None
        return "[Изображение/схема — Gemini Vision не вернул описание]", None
    except Exception as e:
        logger.warning("Vision аннотация изображения не удалась: %s", e)
        return "[Изображение/схема — ошибка аннотации: %s]" % (e,), None


def parse_attachment(
    data: bytes,
    filename: str,
    content_type: Optional[str] = None,
    *,
    skip_images_without_ocr: bool = True,
    context_text: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    Универсальный парсер: PDF (текст + картинки через Vision), DOCX, draw.io (описание через Gemini),
    изображения (только Gemini Vision). context_text — контекст страницы для Gemini.
    skip_images_without_ocr: если True и для изображения Vision недоступен — не ошибка, а короткая пометка.
    """
    ext = _normalize_ext(filename)
    ct = (content_type or "").strip().lower()

    if ext in PDF_EXT or ct in PDF_MIME:
        return extract_text_from_pdf(data, filename, context_text=context_text)
    if ext in DOCX_EXT or ct in DOCX_MIME:
        return extract_text_from_docx(data, filename)
    if ext in XLS_EXT or ct in XLS_MIME:
        return extract_text_from_xls(data, filename)
    if ext in XLSX_EXT or ct in XLSX_MIME:
        return extract_text_from_xlsx(data, filename)
    # Draw.io — по расширению или по содержимому (в Confluence часто без расширения: "Gazcom VNO NMS access")
    if ext in (".drawio", ".drawio.xml") or (ext == ".xml" and "drawio" in (filename or "").lower()):
        return extract_text_from_drawio(data, filename, context_text=context_text)
    # Draw.io без расширения (в Confluence: "Gazcom VNO NMS access"): по mediaType или по содержимому
    if (not ext or ext == ".xml") and (
        "xml" in ct or "octet-stream" in ct or not ct
    ):
        peek = (data[: 2048] if len(data) > 2048 else data).decode("utf-8", errors="ignore")
        if "<mxfile" in peek or "<mxGraphModel" in peek or (peek.lstrip().startswith("<?xml") and "mxfile" in peek):
            return extract_text_from_drawio(data, filename, context_text=context_text)
    if ext in IMAGE_EXT or (ct and ct.startswith("image/")):
        text, err = extract_text_from_image(data, filename, context_text=context_text)
        if err and skip_images_without_ocr:
            return "[Вложение: изображение — задайте GEMINI_API_KEY для аннотации через Gemini Vision]", None
        return text, err

    # Текстовые типы — как есть (если в будущем добавим)
    if ct.startswith("text/"):
        try:
            return data.decode("utf-8", errors="replace").strip(), None
        except Exception:
            pass
    # Последняя попытка: по содержимому draw.io без расширения и без подходящего mediaType
    if len(data) >= 50:
        peek = (data[: 2048] if len(data) > 2048 else data).decode("utf-8", errors="ignore")
        if "<mxfile" in peek or "<mxGraphModel" in peek:
            return extract_text_from_drawio(data, filename, context_text=context_text)
    return "", f"Неподдерживаемый тип вложения: {filename} ({content_type or '?'})"
